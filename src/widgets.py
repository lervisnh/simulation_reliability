from PyQt5 import QtWidgets, QtCore, QtGui
import os
import sys
import docx
import webbrowser

from ui.login_window import Ui_Form as __LoginForm
from ui.main_window import Ui_MainWindow as __MainWindowForm
from ui.passwor_change_window import Ui_Form as __PasswordChangeForm
from ui.user_manage_window import Ui_Form as __UserManageForm
from ui.user_create_window import Ui_Form as __UserCreateForm
from ui.path_config_window import Ui_Form as __PathConfigForm
from ui.model_import_form import Ui_Form as __ModelImportForm
from ui.model_manage_form import Ui_Form as __ModelManageForm
from ui.params_add_form import Ui_Form as __ParamsAddForm
from ui.params_manage_form import Ui_Form as __ParamsManageForm
from ui.simulate_config_form import Ui_Form as __SimConfigForm

from ui.s_analysis_form import Ui_Form as __SAnalysisForm
from ui.u_analysis_form import Ui_Form as __UAnalysisForm
from ui.reliability_form import Ui_Form as __ReliabilityForm
from ui.report_form import Ui_Form as __ReportForm

from utils import AbstractFunction
from db.models import User, SimModel, Param
from db import make_session
from config import get_run_environ, write_run_environ, base_dir
from functools import partial

float_reg = QtCore.QRegExp('^[+-]?([0-9]*[.])?[0-9]+$')


class LoginPage(QtWidgets.QDialog, __LoginForm, AbstractFunction):
    '''登陆界面'''
    def __init__(self, parent=None):
        super().__init__(parent=parent)
        self.setupUi(self)
        self.move_to_center()

        self.setWindowTitle('用户登陆')

        self.btn_login.clicked.connect(self.login_user)  # 登陆按钮
        self.btn_quit.clicked.connect(self.close)  # 退出按钮

        # 注销
        self.edit_username.setText('0000')
        self.edit_password.setText('')

    def login_user(self):
        print('login user')

        username = self.edit_username.text().strip()
        password = self.edit_password.text().strip()
        if not username:
            return self.show_warning_message(message='用户名必须填写', parent=self)
        if not password:
            return self.show_warning_message(message='密码必须填写', parent=self)

        session = make_session()
        user = session.query(User).filter(User.name == username).first()

        if user is None:
            return self.show_warning_message(message='用户不存在', parent=self)

        if user.password != password:
            return self.show_warning_message(message='密码错误', parent=self)
        parent = self.parent()
        parent.refresh_user(user=user)
        parent.temp_path, parent.app_path = get_run_environ()

        self.parent().show()
        if not parent.temp_path or not parent.app_path:
            parent.config_path()
        self.close()


class MainApp(QtWidgets.QMainWindow, __MainWindowForm, AbstractFunction):
    '''主窗，主程序'''
    def __init__(self):
        super().__init__()

        self.user = None

        # todo: 有限元软件 是什么格式 py 还是 exe
        self.app_path = None  # 有限元软件路径
        self.temp_path = None  # 仿真中间文件路径

        self.setupUi(self)
        self.setWindowTitle('仿真软件')
        self.move_to_center()

        self.action_password_reset.triggered.connect(self.user_change_password)  # 密码修改
        self.action_user_manage.triggered.connect(self.user_manage)  # 用户管理，管理员用户
        self.action_edit_config.triggered.connect(self.config_path)  # 配置文件编辑
        self.action_view_config.triggered.connect(self.view_path)  # 配置文件查看

        self.widgets = []  # 存放所有的切换的组件

        grid_layout = QtWidgets.QGridLayout(self.widget_content)  # 右侧区域栅格布局

        self.model_import_widget = ModelImportWidget(parent=self.widget_content, matser=self)  # 导入模型
        grid_layout.addWidget(self.model_import_widget, 0, 0, 1, 1)
        self.widgets.append(self.model_import_widget)

        self.model_manage_widget = ModelManageWidget(master=self, parent=self.widget_content)  # 模型管理
        grid_layout.addWidget(self.model_manage_widget, 0, 0, 1, 1)
        self.widgets.append(self.model_manage_widget)

        self.params_add_widget = ParamsCreateWidget(master=self, parent=self.widget_content)  # 环境参数新建
        grid_layout.addWidget(self.params_add_widget, 0, 0, 1, 1)
        self.widgets.append(self.params_add_widget)

        self.params_manage_widget = ParamsManageWidget(master=self, parent=self.widget_content)  # 环境参数管理
        grid_layout.addWidget(self.params_manage_widget, 0, 0, 1, 1)
        self.widgets.append(self.params_manage_widget)

        self.sim_start_widget = SimConfigWidget(master=self, parent=self.widget_content)  # 开始仿真
        grid_layout.addWidget(self.sim_start_widget, 0, 0, 1, 1)
        self.widgets.append(self.sim_start_widget)

        self.s_analysis_widget = SAnalysisWidget(master=self, parent=self.widget_content)  # 应力应变分析
        grid_layout.addWidget(self.s_analysis_widget, 0, 0, 1, 1)
        self.widgets.append(self.s_analysis_widget)

        self.u_analysis_widget = UAnalysisWidget(master=self, parent=self.widget_content)  # 位移变形分析
        grid_layout.addWidget(self.u_analysis_widget, 0, 0, 1, 1)
        self.widgets.append(self.u_analysis_widget)

        self.reliability_widget = ReliabilityWidget(master=self, parent=self.widget_content)  # 可靠性分析
        grid_layout.addWidget(self.reliability_widget, 0, 0, 1, 1)
        self.widgets.append(self.reliability_widget)

        self.report_widget = ReportWidget(master=self, parent=self.widget_content)  # 生成报告
        grid_layout.addWidget(self.report_widget, 0, 0, 1, 1)
        self.widgets.append(self.report_widget)

        self.btn_add_model.clicked.connect(partial(self.switch_to, name='model_import'))  # 导入模型按钮
        self.btn_delete_model.clicked.connect(partial(self.switch_to, name='model_manage'))  # 删除模型按钮
        self.btn_add_paras.clicked.connect(partial(self.switch_to, name='params_add'))  # 新增参数 按钮
        self.btn_delete_paras.clicked.connect(partial(self.switch_to, name='params_manage'))  # 删除参数按钮
        self.btn_simulation.clicked.connect(partial(self.switch_to, name='sim_start'))  # 开始仿真按钮
        self.btn_stress_analysis.clicked.connect(partial(self.switch_to, name='s_analysis'))  # 应力应变分析
        self.btn_displacement_analysis.clicked.connect(partial(self.switch_to, name='u_analysis'))  # 位移变形分析
        self.btn_reliability_analysis.clicked.connect(partial(self.switch_to, name='reliability_analysis'))  # 可靠性分析
        self.btn_report.clicked.connect(partial(self.switch_to, name='report'))  # 报告

        self.btn_simulation.click()

    def refresh_user(self, user: User):
        '''更新主程序的 user 属性'''
        if not isinstance(user, User):
            raise TypeError
        self.user = user
        if self.user.is_admin is False:  # 管理员用户才有 用户管理
            self.action_user_manage.setDisabled(True)
        else:
            self.action_user_manage.setDisabled(False)

    def user_change_password(self):
        '''弹窗 用户修改密码 '''
        d = ChangeUserPasswordDialog(user=self.user, parent=self)
        d.show()

    def user_manage(self):
        '''弹窗 用户管理'''
        d = UserManagePage(user=self.user, parent=self)
        d.show()

    def config_path(self):
        '''弹窗 配置路径 '''
        d = PathConfigPage(app_path=self.app_path, temp_path=self.temp_path, parent=self)
        d.show()

    def view_path(self):
        '''弹窗 查看配置 '''
        d = PathViewPage(app_path=self.app_path, temp_path=self.temp_path, parent=self)
        d.show()

    def switch_to(self, name):
        '''切换 右侧的功能 '''
        for w in self.widgets:
            w.hide()  # 隐藏所有的组件
        if name == 'model_import':
            self.model_import_widget.show()
        elif name == 'model_manage':
            self.model_manage_widget.show()
            self.model_manage_widget.refresh_table()

        elif name == 'params_add':
            self.params_add_widget.show()

        elif name == 'params_manage':
            self.params_manage_widget.show()
            self.params_manage_widget.refresh_table()

        elif name == 'sim_start':
            self.sim_start_widget.show()
            self.sim_start_widget.refresh_table_model()
            self.sim_start_widget.refresh_table_params()

        elif name == 's_analysis':
            self.s_analysis_widget.show()
            self.s_analysis_widget.refresh_table()

        elif name == 'u_analysis':
            self.u_analysis_widget.show()
            self.u_analysis_widget.refresh_table()

        elif name == 'reliability_analysis':
            self.reliability_widget.show()
            self.reliability_widget.refresh_table()

        elif name == 'report':
            self.report_widget.show()
            self.report_widget.refresh_table()

class ChangeUserPasswordDialog(QtWidgets.QDialog, __PasswordChangeForm, AbstractFunction):
    '''用户修改密码 '''
    def __init__(self, user, parent=None):
        super().__init__(parent=parent)
        self.setupUi(self)
        if not isinstance(user, User):
            raise TypeError
        self.user = user

        self.btn_quit.clicked.connect(self.close)
        self.btn_confirm.clicked.connect(self.change_password)

    def change_password(self):
        old_password = self.edit_old_password.text().strip()
        new_password = self.edit_new_password.text().strip()

        if not old_password:
            return self.show_warning_message(message='旧密码必须填写', parent=self.parent())
        if not new_password:
            return self.show_warning_message(message='新密码必须填写', parent=self.parent())

        if old_password != self.user.password:
            return self.show_warning_message(message='旧密码验证失败', parent=self.parent())

        if new_password == old_password:
            return
        s = make_session()
        user = s.query(User).filter(User.id == self.user.id).first()
        if user is None:
            raise ValueError
        user.password = new_password
        s.commit()
        self.parent().refresh_user(user=user)
        self.show_warning_message(message='密码修改成功', parent=self.parent())
        self.close()


class UserManagePage(QtWidgets.QDialog, __UserManageForm, AbstractFunction):
    '''用户管理 '''
    def __init__(self, user, parent=None):
        super().__init__(parent=parent)
        if not isinstance(user, User):
            raise TypeError
        self.user = user
        self.setupUi(self)

        self.setWindowTitle('用户管理')

        self.refresh_table()
        self.btn_create.clicked.connect(self.create_new_user)
        self.btn_reset.clicked.connect(self.reset_user_password)
        self.btn_delete.clicked.connect(self.delete_one_user)

        self.table_users.contextMenuEvent = self.table_right_click

    def refresh_table(self):
        # self.table_users.clear()
        self.clear_table()

        s = make_session()
        users = s.query(User).all()
        rows = len(users)
        columns = 3
        self.table_users.setRowCount(rows)
        self.table_users.setColumnCount(columns)

        r = rows - 1
        for user in users:
            self.table_users.setItem(r, 0, QtWidgets.QTableWidgetItem(str(user.id)))
            self.table_users.setItem(r, 1, QtWidgets.QTableWidgetItem(str(user.name)))
            print(user.name)
            self.table_users.setItem(r, 2, QtWidgets.QTableWidgetItem(str(
                '是' if user.is_admin else '否')))
            r -= 1

    def clear_table(self):
        for i in range(self.table_users.rowCount()):
            self.table_users.removeRow(0)

    def create_new_user(self):
        ''''''
        print('create_new_user')
        d = UserCreatePage(user=self.user, parent=self)
        d.show()

    def table_right_click(self, a0):
        context_menu = QtWidgets.QMenu(self)

        new_action = QtWidgets.QAction('新建用户', context_menu)
        reset_action = QtWidgets.QAction('重置密码', context_menu)
        delete_action = QtWidgets.QAction('删除用户', context_menu)
        context_menu.addAction(new_action)
        context_menu.addAction(reset_action)
        context_menu.addAction(delete_action)

        new_action.triggered.connect(self.create_new_user)
        reset_action.triggered.connect(self.reset_user_password)
        delete_action.triggered.connect(self.delete_one_user)

        item = self.table_users.currentItem()
        if not isinstance(item, QtWidgets.QTableWidgetItem):
            return

        context_menu.show()
        context_menu.exec_(QtGui.QCursor.pos())

    def reset_user_password(self):
        item = self.table_users.currentItem()
        row = self.table_users.currentRow()
        _id_item = self.table_users.item(row, 0)
        _id = int(_id_item.text())
        name = self.table_users.item(row, 1).text()
        self.table_users.setCurrentItem(_id_item)

        if _id == self.user.id:
            return self.show_warning_message(message='请使用“修改密码”功能', parent=self)
        reply = self.show_warning_message(message='确定重置该用户的密码吗？<br>id：{}<br>用户名: {}'.format(_id, name), parent=self)
        if reply is False:
            return
        print(item, row, _id_item.text())
        s = make_session()
        user = s.query(User).filter(User.id == _id).first()
        if user is None:
            return self.refresh_table()
        user.password = '123456'
        s.commit()
        return self.show_warning_message(message='用户密码重置成功<br>id：{}<br>用户名: {}<br>密码: 123456'.format(_id, name), parent=self)

    def delete_one_user(self):
        item = self.table_users.currentItem()
        row = self.table_users.currentRow()
        _id_item = self.table_users.item(row, 0)
        _id = int(_id_item.text())
        name = self.table_users.item(row, 1).text()
        self.table_users.setCurrentItem(_id_item)

        if _id == self.user.id:
            return self.show_warning_message(message='非法删除操作', parent=self)
        reply = self.show_warning_message(message='确定删除用户吗？<br>id：{}<br>用户名: {}'.format(_id, name), parent=self)
        if reply is False:
            return
        print(item, row, _id_item.text())
        s = make_session()
        user = s.query(User).filter(User.id == _id).first()
        if user is None:
            return self.refresh_table()

        s.delete(user)
        s.commit()
        return self.refresh_table()


class UserCreatePage(QtWidgets.QDialog, __UserCreateForm, AbstractFunction):
    '''新建用户'''
    def __init__(self, user, parent=None):
        if not isinstance(user, User):
            raise TypeError
        super().__init__(parent=parent)
        self.setupUi(self)

        self.btn_confirm.clicked.connect(self.add_new_user)
        self.btn_quit.clicked.connect(self.close)

    def add_new_user(self):
        name = self.lineEdit.text().strip()
        if not name:
            return self.show_warning_message(message='用户名必须填写', parent=self.parent())
        s = make_session()
        user = s.query(User).filter(User.name == name).first()
        if user is not None:
            return self.show_warning_message(message='用户名“{}”已经被占用'.format(name), parent=self.parent())

        user = User(name=name, password='1234')
        s.add(user)
        s.commit()
        self.parent().refresh_table()
        self.show_warning_message(message='用户添加成功, 初始密码是 1234', parent=self.parent())
        self.close()


class PathConfigPage(QtWidgets.QDialog, __PathConfigForm, AbstractFunction):
    '''仿真软件配置'''
    def __init__(self, app_path, temp_path, parent=None):
        super().__init__(parent=parent)
        self.app_path = app_path
        self.temp_path = temp_path
        self.setupUi(self)
        if app_path is None:
            app_path = ''
        if temp_path is None:
            temp_path = ''

        self.setWindowTitle('配置仿真路径')

        self.text_sim_middle.setPlainText(temp_path)  # 仿真中间文件路径
        self.text_app_path.setPlainText(app_path)  # 有限元软件路径

        self.btn_sim_look.clicked.connect(self.ask_path)
        self.btn_app_look.clicked.connect(self.ask_exe_path)
        self.btn_config.clicked.connect(self.save_config)
        self.btn_quit.clicked.connect(self.close)

    def ask_path(self):
        # fp, _ = QtWidgets.QFileDialog.getOpenFileName(
        fp = QtWidgets.QFileDialog.getExistingDirectory(
            parent=self,
            caption='仿真中间文件',
            directory='.',
            # filter='JNL (*.jnl)'
        )
        # todo: 仿真中间文件 的格式是 jnl 还是 py 或者其他，作相应修改
        if not fp:
            return
        self.text_sim_middle.setPlainText(fp)
        self.temp_path = fp

    def ask_exe_path(self):
        fp, _ = QtWidgets.QFileDialog.getOpenFileName(
            parent=self,
            caption='有限元软件路径',
            directory='.',
            filter='有限元软件 (*.*)'
        )
        # todo: 有限元软件路径 的格式是 exe 还是 py 或者其他，作相应修改
        print(fp)
        if not fp:
            return
        self.text_app_path.setPlainText(fp)
        self.app_path = fp

    def save_config(self):
        if not self.app_path:
            return self.show_warning_message(message='', parent=self)
        if not self.temp_path:
            return self.show_warning_message(message='', parent=self)
        write_run_environ(temp_path=self.temp_path, app_path=self.app_path)
        parent = self.parent()
        parent.app_path = self.app_path
        parent.temp_path = self.temp_path
        self.close()


class PathViewPage(PathConfigPage):
    '''仿真软件配置 查看 '''
    def __init__(self, app_path, temp_path, parent=None):
        super().__init__(app_path, temp_path, parent)
        self.btn_app_look.hide()
        self.btn_sim_look.hide()
        self.setWindowTitle('查看仿真路径')

        self.btn_config.clicked.disconnect(self.save_config)
        self.btn_config.clicked.connect(self.close)

        self.label_header.setText('查看仿真路径')


class ModelImportWidget(QtWidgets.QWidget, __ModelImportForm, AbstractFunction):
    '''模型导入 '''
    def __init__(self, matser, parent=None):
        super().__init__(parent)
        self.setupUi(self)

        if not isinstance(matser, MainApp):
            raise TypeError
        self.master = matser

        self.temp_path = None
        self.temp_ext_path = None
        self.btn_temp_ext_load.setDisabled(True)

        self.check_btn.clicked.connect(self.toggle_ext)
        self.btn_temp_load.clicked.connect(self.ask_jnl_path)
        self.btn_temp_ext_load.clicked.connect(self.ask_ext_path)
        self.btn_confirm.clicked.connect(self.do_import)

    def toggle_ext(self):
        if self.check_btn.isChecked():
            self.btn_temp_ext_load.setDisabled(False)
        else:
            self.btn_temp_ext_load.setDisabled(True)

    def ask_jnl_path(self):
        fp, _ = QtWidgets.QFileDialog.getOpenFileName(
            parent=self,
            caption='模型文件',
            directory='.',
            filter='JNL (*.jnl)'
        )
        # todo: 模型文件 的格式是 jnl 还是 py 或者其他，作相应修改
        if not fp:
            return
        self.edit_jnl_path.setText(fp)
        self.temp_path = fp

    def ask_ext_path(self):
        '''三维模型非仿真软件生成'''
        fp, _ = QtWidgets.QFileDialog.getOpenFileName(
            parent=self,
            caption='三维模型',
            directory='.',
            filter='JNL (*.*)'
        )
        # todo: 三维模型路径 的格式是 jnl 还是 py 或者其他，作相应修改
        if not fp:
            return
        self.edit_external_model_path.setText(fp)
        self.temp_ext_path = fp

    def do_import(self):
        if self.temp_path is None:
            return self.show_warning_message(message='模型文件必须正确载入', parent=self)

        user = self.master.user
        _, jnl_name = os.path.split(self.temp_path)
        if self.temp_ext_path is not None:
            _, ext_name = os.path.split(self.temp_ext_path)
        else:
            ext_name = None
        s = make_session()

        sim = SimModel(
            name=jnl_name, jnl_path=jnl_name,
            external_model_path=ext_name, user_id=user.id)
        s.add(sim)
        s.commit()
        dir_path = os.path.join(base_dir, 'data', 'models')
        if not os.path.exists(dir_path):
            os.makedirs(dir_path, exist_ok=True)
        _, __ext = os.path.splitext(jnl_name)
        __ext = __ext[1:]

        sim.jnl_path = '{}.{}'.format(sim.id, __ext)

        with open(os.path.join(dir_path, sim.jnl_path), 'wb') as f:
            with open(self.temp_path, 'rb')  as g:
                for line in g:
                    f.write(line)

        if self.temp_ext_path:
            _, __ext = os.path.splitext(ext_name)
            __ext = __ext[1:]
            sim.external_model_path = '{}-ext.{}'.format(sim.id, __ext)

            with open(os.path.join(dir_path, sim.external_model_path), 'wb') as f:
                with open(self.temp_ext_path, 'rb')  as g:
                    for line in g:
                        f.write(line)

        s.commit()

        self.re_init()
        self.show_warning_message(message='模型 {} 添加成功'.format(sim.id), parent=self)

    def re_init(self):
        self.edit_external_model_path.setText('')
        self.edit_jnl_path.setText('')
        self.temp_ext_path = None
        self.temp_path = None

        self.check_btn.setChecked(False)


class ModelManageWidget(QtWidgets.QWidget, __ModelManageForm, AbstractFunction):
    '''模型管理删除'''
    def __init__(self, master, parent=None):
        super().__init__(parent=parent)

        if not isinstance(master, MainApp):
            raise TypeError
        self.master = master

        self.setupUi(self)
        self.edit_id_to_del.setText('')
        self.edit_id_to_del.setReadOnly(True)

        self.refresh_table()
        self.table_content.clicked.connect(self.chose_one)
        self.btn_delete.clicked.connect(self.delete_one)
        self.table_content.contextMenuEvent = self.table_right_click

    def refresh_table(self):

        self.clear_table()

        user = self.master.user
        if user is None:
            return

        s = make_session()
        sim_ls = s.query(SimModel).filter(SimModel.user_id == user.id).all()
        rows = len(sim_ls)
        columns = 4

        self.table_content.setRowCount(rows)
        self.table_content.setColumnCount(columns)

        r = rows - 1
        for sim in sim_ls:

            item = QtWidgets.QTableWidgetItem(str(sim.id))
            item.setTextAlignment(QtCore.Qt.AlignCenter)
            self.table_content.setItem(r, 0, item)

            item = QtWidgets.QTableWidgetItem(str('' if sim.name is None else sim.name))
            item.setTextAlignment(QtCore.Qt.AlignCenter)
            self.table_content.setItem(r, 1, item)

            item = QtWidgets.QTableWidgetItem(str('' if sim.jnl_path is None else sim.jnl_path))
            item.setTextAlignment(QtCore.Qt.AlignCenter)
            self.table_content.setItem(r, 2, item)

            item = QtWidgets.QTableWidgetItem(str('' if sim.external_model_path is None else sim.external_model_path))
            item.setTextAlignment(QtCore.Qt.AlignCenter)
            self.table_content.setItem(r, 3, item)

            r -= 1
        self.edit_id_to_del.setText('')

    def clear_table(self):
        for i in range(self.table_content.rowCount()):
            self.table_content.removeRow(0)

    def chose_current(self):
        row = self.table_content.currentRow()
        item = self.table_content.item(row, 0)
        if not isinstance(item, QtWidgets.QTableWidgetItem):
            return
        self.edit_id_to_del.setText(item.text())

    def chose_one(self):
        self.chose_current()

    def delete_one(self):
        s = make_session()
        _id = self.edit_id_to_del.text().strip()
        try:
            _id = int(_id)
        except (ValueError, TypeError):
            _id = None
        if not _id:
            return self.show_warning_message(message='请先选择一个模型', parent=self)

        sim = s.query(SimModel).filter(SimModel.id == _id).first()
        if sim is None:
            return self.refresh_table()
        s.delete(sim)
        s.commit()

        # 删除文件
        fp = os.path.join(base_dir, 'data', 'models', sim.jnl_path)
        if os.path.exists(fp):
            os.remove(fp)
        if sim.external_model_path:
            fp = os.path.join(base_dir, 'data', 'models', sim.external_model_path)
            if os.path.exists(fp):
                os.remove(fp)

        self.refresh_table()

    def table_right_click(self, a0):
        context_menu = QtWidgets.QMenu(self)
        delete_action = QtWidgets.QAction('删除模型', context_menu)

        context_menu.addAction(delete_action)

        delete_action.triggered.connect(self.delete_one)

        item = self.table_content.currentItem()
        if not isinstance(item, QtWidgets.QTableWidgetItem):
            return
        self.chose_current()

        context_menu.show()
        context_menu.exec_(QtGui.QCursor.pos())


class ParamsCreateWidget(QtWidgets.QWidget, __ParamsAddForm, AbstractFunction):
    '''新建参数'''
    def __init__(self, master, parent=None):
        super().__init__(parent=parent)
        self.setupUi(self)

        if not isinstance(master, MainApp):
            raise TypeError

        self.master = master
        self.batch_set_float_validate()
        self.btn_confirm.clicked.connect(self.add_one)

        self.params_needs = {
            "均匀分布": ['location', 'size'],
            "指数分布": ['size'],
            "正态分布": ['location', 'size'],
            "对数正态分布": ['size', 'shape'],
            "威布尔分布": ['size', 'shape'],
            "固定值": ['location'],
        }

        self.optional_params_map = {
            'location': '位置参数',
            'size': '尺度参数',
            'shape': '形状参数'
        }

        self.combo_distr_kinds.clear()
        self.combo_distr_kinds.currentIndexChanged.connect(self.re_config)
        self.combo_distr_kinds.addItems([
            "均匀分布",
            "指数分布",
            "正态分布",
            "对数正态分布",
            "威布尔分布",
            "固定值"
        ])

        self.combo_distr_kinds.setCurrentIndex(0)

    def re_config(self):
        kind = self.combo_distr_kinds.currentText()
        for w in [self.edit_shape, self.edit_size, self.edit_location]:
            w.setDisabled(False)

        __all = ['location', 'size', 'shape']
        for k in set(__all) - set(self.params_needs[kind]):
            w = self.__dict__['edit_{}'.format(k)]
            w.setDisabled(True)
            w.setText('')

    def batch_set_float_validate(self):
        for k, v in self.__dict__.items():
            if k in ['edit_shape', 'edit_location', 'edit_size'] and isinstance(v, QtWidgets.QLineEdit):
                validator = QtGui.QRegExpValidator(float_reg, v)  # 创建一个浮点型数据验证器
                v.setValidator(validator)  # 将验证器与输入框绑定

    def add_one(self):
        name = self.edit_name.text().strip()
        kind = self.combo_distr_kinds.currentText()
        location = self.edit_location.text().strip()
        try:
            location = float(location)
        except (TypeError, ValueError):
            location = None

        size = self.edit_size.text().strip()
        try:
            size = float(size)
        except (TypeError, ValueError):
            size = None
        shape = self.edit_shape.text().strip()
        try:
            shape = float(shape)
        except (TypeError, ValueError):
            shape = None
        unit = self.edit_unit.text().strip()

        if not name:
            return self.show_warning_message(message='参数名称必须填写', parent=self)
        if not kind:
            return self.show_warning_message(message='分布类型必须填写', parent=self)
        if not unit:
            return self.show_warning_message(message='单位必须填写', parent=self)

        for k in self.params_needs[kind]:
            if not isinstance(locals().get(k), float):
                return self.show_warning_message(message='{}必须填写'.format(self.optional_params_map[k]), parent=self)

        s = make_session()
        p = Param(
            name=name, distri=kind,
            location_para=location, scale_para=size,
            shape_para=shape,
            unit=unit,
            user_id=self.master.user.id
        )
        s.add(p)
        s.commit()

        for k, v in self.__dict__.items():
            if isinstance(v, QtWidgets.QLineEdit):
                v.setText('')
        self.combo_distr_kinds.setCurrentIndex(0)
        self.show_warning_message(message='参数"{}"添加成功'.format(name), parent=self)


class ParamsManageWidget(QtWidgets.QWidget, __ParamsManageForm, AbstractFunction):
    '''参数管理'''
    def __init__(self, master, parent=None):
        super().__init__(parent=parent)
        self.setupUi(self)
        self.master = master
        self.edit_id.setText('')
        self.edit_id.setReadOnly(True)

        self.btn_del.clicked.connect(self.delete_one)
        self.table_content.clicked.connect(self.chose_one)
        self.table_content.contextMenuEvent = self.table_right_click

    def refresh_table(self):

        self.clear_table()
        user = self.master.user
        if user is None:
            return

        s = make_session()
        param_ls = s.query(Param).filter(Param.user_id == user.id).all()
        rows = len(param_ls)
        columns = 7
        print(rows, columns)

        self.table_content.setRowCount(rows)
        self.table_content.setColumnCount(columns)

        r = rows - 1
        for p in param_ls:
            item = QtWidgets.QTableWidgetItem(str(p.id))
            item.setTextAlignment(QtCore.Qt.AlignCenter)
            self.table_content.setItem(r, 0, item)

            item = QtWidgets.QTableWidgetItem(str('' if p.name is None else p.name))
            item.setTextAlignment(QtCore.Qt.AlignCenter)
            self.table_content.setItem(r, 1, item)

            item = QtWidgets.QTableWidgetItem(str('' if p.distri is None else p.distri))
            item.setTextAlignment(QtCore.Qt.AlignCenter)
            self.table_content.setItem(r, 2, item)

            item = QtWidgets.QTableWidgetItem(str('' if p.location_para is None else p.location_para))
            item.setTextAlignment(QtCore.Qt.AlignCenter)
            self.table_content.setItem(r, 3, item)

            item = QtWidgets.QTableWidgetItem(str('' if p.shape_para is None else p.shape_para))
            item.setTextAlignment(QtCore.Qt.AlignCenter)
            self.table_content.setItem(r, 4, item)

            item = QtWidgets.QTableWidgetItem(str('' if p.scale_para is None else p.scale_para))
            item.setTextAlignment(QtCore.Qt.AlignCenter)
            self.table_content.setItem(r, 5, item)

            item = QtWidgets.QTableWidgetItem(str('' if p.unit is None else p.unit))
            item.setTextAlignment(QtCore.Qt.AlignCenter)
            self.table_content.setItem(r, 6, item)

            r -= 1

        self.edit_id.setText('')

    def clear_table(self):
        for i in range(self.table_content.rowCount()):
            self.table_content.removeRow(0)

    def chose_current(self):
        row = self.table_content.currentRow()
        item = self.table_content.item(row, 0)
        if not isinstance(item, QtWidgets.QTableWidgetItem):
            return
        self.edit_id.setText(item.text())

    def chose_one(self):
        self.chose_current()

    def delete_one(self):
        s = make_session()
        _id = self.edit_id.text().strip()
        try:
            _id = int(_id)
        except (ValueError, TypeError):
            _id = None
        if not _id:
            return self.show_warning_message(message='请先选择一个模型', parent=self)

        sim = s.query(Param).filter(Param.id == _id).first()
        if sim is None:
            return self.refresh_table()
        s.delete(sim)
        s.commit()

        self.refresh_table()

    def table_right_click(self, a0):
        context_menu = QtWidgets.QMenu(self)
        delete_action = QtWidgets.QAction('删除参数', context_menu)

        context_menu.addAction(delete_action)

        delete_action.triggered.connect(self.delete_one)

        item = self.table_content.currentItem()
        if not isinstance(item, QtWidgets.QTableWidgetItem):
            return
        self.chose_current()

        context_menu.show()
        context_menu.exec_(QtGui.QCursor.pos())


class SimConfigWidget(QtWidgets.QWidget, __SimConfigForm, AbstractFunction):
    '''开始仿真'''
    def __init__(self, master, parent=None):
        if not isinstance(master, MainApp):
            raise TypeError
        super().__init__(parent=parent)
        self.setupUi(self)
        if not isinstance(master, MainApp):
            raise TypeError
        self.master = master

        self.edit_model.setText('')
        self.edit_params.setText('')

        self.table_params.clicked.connect(self.chose_current_params)
        self.table_model.clicked.connect(self.chose_current_model)
        self.table_model.contextMenuEvent = self.table_model_right_click
        self.table_params.contextMenuEvent = self.table_params_right_click

        self.btn_confirm.clicked.connect(self.confirm_params_and_model)
        self.btn_start.clicked.connect(self.run_simulation)

    def refresh_table_model(self):

        self.clear_table_model()
        user = self.master.user
        if user is None:
            return

        s = make_session()
        sim_ls = s.query(SimModel).filter(SimModel.user_id == user.id).all()
        rows = len(sim_ls)
        columns = 4

        self.table_model.setRowCount(rows)
        self.table_model.setColumnCount(columns)

        r = rows - 1
        for sim in sim_ls:
            item = QtWidgets.QTableWidgetItem(str(sim.id))
            item.setTextAlignment(QtCore.Qt.AlignCenter)
            self.table_model.setItem(r, 0, item)

            item = QtWidgets.QTableWidgetItem(str('' if sim.name is None else sim.name))
            item.setTextAlignment(QtCore.Qt.AlignCenter)
            self.table_model.setItem(r, 1, item)

            item = QtWidgets.QTableWidgetItem(str('' if sim.jnl_path is None else sim.jnl_path))
            item.setTextAlignment(QtCore.Qt.AlignCenter)
            self.table_model.setItem(r, 2, item)

            item = QtWidgets.QTableWidgetItem(str('' if sim.external_model_path is None else sim.external_model_path))
            item.setTextAlignment(QtCore.Qt.AlignCenter)
            self.table_model.setItem(r, 3, item)

            r -= 1
        self.edit_model.setText('')

    def clear_table_model(self):
        for i in range(self.table_model.rowCount()):
            self.table_model.removeRow(0)

    def refresh_table_params(self):

        self.clear_table_params()
        user = self.master.user
        if user is None:
            return

        s = make_session()
        param_ls = s.query(Param).filter(Param.user_id == user.id).all()
        rows = len(param_ls)
        columns = 7
        print(rows, columns)

        self.table_params.setRowCount(rows)
        self.table_params.setColumnCount(columns)

        r = rows - 1
        for p in param_ls:
            item = QtWidgets.QTableWidgetItem(str(p.id))
            item.setTextAlignment(QtCore.Qt.AlignCenter)
            self.table_params.setItem(r, 0, item)

            item = QtWidgets.QTableWidgetItem(str('' if p.name is None else p.name))
            item.setTextAlignment(QtCore.Qt.AlignCenter)
            self.table_params.setItem(r, 1, item)

            item = QtWidgets.QTableWidgetItem(str('' if p.distri is None else p.distri))
            item.setTextAlignment(QtCore.Qt.AlignCenter)
            self.table_params.setItem(r, 2, item)

            item = QtWidgets.QTableWidgetItem(str('' if p.location_para is None else p.location_para))
            item.setTextAlignment(QtCore.Qt.AlignCenter)
            self.table_params.setItem(r, 3, item)

            item = QtWidgets.QTableWidgetItem(str('' if p.shape_para is None else p.shape_para))
            item.setTextAlignment(QtCore.Qt.AlignCenter)
            self.table_params.setItem(r, 4, item)

            item = QtWidgets.QTableWidgetItem(str('' if p.scale_para is None else p.scale_para))
            item.setTextAlignment(QtCore.Qt.AlignCenter)
            self.table_params.setItem(r, 5, item)

            item = QtWidgets.QTableWidgetItem(str('' if p.unit is None else p.unit))
            item.setTextAlignment(QtCore.Qt.AlignCenter)
            self.table_params.setItem(r, 6, item)

            r -= 1

        self.edit_params.setText('')

    def clear_table_params(self):
        for i in range(self.table_params.rowCount()):
            self.table_params.removeRow(0)

    def chose_current_model(self, a0=None):
        row = self.table_model.currentRow()
        item = self.table_model.item(row, 0)
        if not isinstance(item, QtWidgets.QTableWidgetItem):
            return
        self.edit_model.setText(item.text())

    def chose_one_model(self):
        self.chose_current_model()

    def chose_current_params(self, a0=None):
        row = self.table_params.currentRow()
        item = self.table_params.item(row, 0)
        if not isinstance(item, QtWidgets.QTableWidgetItem):
            return
        self.edit_params.setText(item.text())

    def chose_one_params(self):
        self.chose_current_params()

    def table_model_right_click(self, a0):
        context_menu = QtWidgets.QMenu(self)
        delete_action = QtWidgets.QAction('选择模型', context_menu)

        context_menu.addAction(delete_action)

        delete_action.triggered.connect(self.chose_current_model)

        item = self.table_model.currentItem()
        if not isinstance(item, QtWidgets.QTableWidgetItem):
            return

        context_menu.show()
        context_menu.exec_(QtGui.QCursor.pos())

    def table_params_right_click(self, a0):
        context_menu = QtWidgets.QMenu(self)
        delete_action = QtWidgets.QAction('选择参数', context_menu)

        context_menu.addAction(delete_action)

        delete_action.triggered.connect(self.chose_current_params)

        item = self.table_params.currentItem()
        if not isinstance(item, QtWidgets.QTableWidgetItem):
            return

        context_menu.show()
        context_menu.exec_(QtGui.QCursor.pos())

    def confirm_params_and_model(self):
        '''确认选择'''
        model_id = self.edit_params.text().strip()
        params_id = self.edit_params.text().strip()

        if not model_id:
            return self.show_warning_message(message='请先选择一个模型', parent=self)
        if not params_id:
            return self.show_warning_message(message='请先选择一组参数', parent=self)

        return self.show_warning_message(message='请点击 “开始仿真”', parent=self)

    def run_simulation(self):
        '''执行模拟仿真'''

        if not self.check_config_environ_path():
            return 
        app_path = self.master.app_path  # 有限元软件路径
        temp_path = self.master.temp_path  # 中间件路径

        model_id = self.edit_model.text().strip()
        params_id = self.edit_params.text().strip()

        try:
            model_id = int(model_id)
        except (TypeError, ValueError):
            model_id = None

        try:
            params_id = int(params_id)
        except (TypeError, ValueError):
            params_id = None

        if not model_id:
            return self.show_warning_message(message='请先选择一个模型', parent=self)
        if not params_id:
            return self.show_warning_message(message='请先选择一组参数', parent=self)

        s = make_session()

        model = s.query(SimModel).filter(SimModel.id == model_id).first()
        params = s.query(Param).filter(Param.id == params_id).first()

        if not isinstance(model, SimModel):
            raise ValueError
        if not isinstance(params, Param):
            raise ValueError

        jnl_path = os.path.join(base_dir, 'data', 'models', model.jnl_path)  # 模型由abaqus生成的jnl文件
        if isinstance(model.external_model_path, str):
            ext_path = os.path.join(base_dir, 'data', 'models', model.external_model_path)
        else:
            ext_path = None  # 如果jnl文件的模型时外部导入的，需要其路径

        mapping = { "均匀分布":"uniform_distri",
                    "指数分布":"exponential_distri",
                    "正态分布":"normal_distri",
                    "对数正态分布":"lognormal_distri",
                    "威布尔分布":"weibull_distri",
                    "固定值":"figure_no_distri" }
        name = params.name  # 载荷类型(名称)
        distri = mapping[params.distri]  # 载荷服从分布
        location_para = params.location_para  # 位置参数
        scale_para = params.scale_para  # 比例参数
        shape_para = params.shape_para  # 形状参数
        unit = params.unit  # 载荷的单位

        from functions.generate_randoms import generate_randoms as __RandomsGenerator
        from functions.copy_jnl_modify_paras import copy_jnl_modify_paras as __copy_jnl_modify_paras

        if distri=="figure_no_distri":
            randoms = [location_para]*10
        else:
            randoms = __RandomsGenerator(distri=distri, 
                                         location_para=location_para, 
                                         scale_para=scale_para, 
                                         shape_para=shape_para, 
                                         number_rand=10 ).generating()


        content = __copy_jnl_modify_paras(copy_path=jnl_path, ext_path=ext_path, to_add_paras=randoms[0])
        content.write_to_file( str(model_id) )

        from functions.cmd_file import cmd_file
        self.show_warning_message(message='确认开始仿真', parent=self)
        cmd_file( str(model_id) )
        job_name = content._job_name()
        if os.path.exists( temp_path+"/model_"+str(model_id)+"/"+job_name+".odb" ):
            # 输出数据到txt
            u_txt_path = temp_path+'\\'+'model_'+str(model_id)+'\\'+job_name+'_U.txt' # 位移变形数据
            s_txt_path = temp_path+'\\'+'model_'+str(model_id)+'\\'+job_name+'_S.txt' # 应力应变数据
            if not os.path.exists( u_txt_path ) or not os.path.exists( s_txt_path ):
                from functions.abaqus_postprocesser import abaqus_postprocesser as __ABAQUS_PostProcesser
                odb_file_path = self.master.temp_path+'\\'+'model_'+str(model_id)+'\\'+job_name+'.odb'
                __ABAQUS_PostProcesser(odb_path = odb_file_path).reading_data()
            # 初始创建报告, 写入模型文件名、所选参数等
            self.init_report(model=model, params=params)
            self.show_warning_message(message='仿真已完成', parent=self)
        else:
            self.show_warning_message(message='警告: 仿真中止!\n 请检查以下内容:\n1. ABAQUS服务是否已开启\n2. 仿真配置或模型文件仿真设置(如网格设置、边界条件设置等)!', parent=self)


        
        #print('run simulation')
        #print(locals())

    
    def check_config_environ_path(self):
        '''检查有限元软件路径 和 中间件路径'''
        if os.path.exists(self.master.app_path) and os.path.exists(self.master.temp_path):
            return True
        
        if not os.path.exists(self.master.app_path): # 有限元软件路径
            self.show_warning_message(message='有限元软件配置错误，请重新配置!', parent=self)
        
        if not os.path.exists(self.master.temp_path): # 中间件路径
            self.show_warning_message(message='中间文件存储路径不存在，请重新配置!', parent=self)
        
        return False

    # 初始化报告
    def init_report(self, model, params):
        curr_path = os.getcwd() #当前路径
        #！！！注意切换路径！！！
        os.chdir( self.master.temp_path+"/model_"+str(model.id) ) #切换路径

        word_name = 'model_'+str(model.id)+'_report'#文档名称

        def ADD_HEADDINGS(document, contents, level, pt, CENTER=False):
            # level级标题 
            heading = document.add_heading('', level)
            if CENTER: #标题居中
                heading.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            #添加标题
            heading_run = heading.add_run( contents )
            #设置标题
            heading_run.font.size = docx.shared.Pt(pt) #标题字体大小
            heading_run.font.name = u'Times New Roman' #标题字体类型
            r = heading_run._element
            r.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), u'黑体')
            heading_run.font.color.rgb = docx.shared.RGBColor(0,0,0) #黑色
            heading_run.bold = True #加粗

        def ADD_PARAPRAPH(document, contents):
            paragraph = document.add_paragraph('')
            paragraph.paragraph_format.line_spacing = 1.25
            paragraph.paragraph_format.space_after = 0 #段后间距
            paragraph.paragraph_format.space_before = 0 #段前间距
            paragraph = paragraph.add_run( contents )
            paragraph.font.size = docx.shared.Pt(14) #标题字体大小
            paragraph.font.name = u'Times New Roman' #标题字体类型
            r = paragraph._element
            r.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), u'宋体')
            paragraph.font.color.rgb = docx.shared.RGBColor(0,0,0) #黑色
            paragraph.bold = False #加粗
            
            
        document = docx.Document()
        ##############################################################################
        # 一级标题
        ADD_HEADDINGS( document, u'有限元仿真及可靠性分析简报', 0, 22, True )
        ##############################################################################
        #二级标题  模型概述
        ADD_HEADDINGS( document, u'1. 模型概述', 1, 14 )
        '''模型说明'''
        ADD_PARAPRAPH(document, u'模型编号：'+str(model.id))
        ADD_PARAPRAPH(document, u'模型文件名称：'+model.name.split('.')[0])
        ##############################################################################
        #二级标题  环境参数概述
        ADD_HEADDINGS( document, u'2. 环境参数概述', 1, 14 )
        '''添加环境参数说明'''
        ADD_PARAPRAPH(document, u'名称：'+params.name)
        ADD_PARAPRAPH(document, u'服从分布：'+params.distri)
        ADD_PARAPRAPH(document, u'分布的位置参数(或固定值数值)：'+str(params.location_para))  # 位置参数
        ADD_PARAPRAPH(document, u'分布的比例参数：'+str(params.scale_para))  # 比例参数
        ADD_PARAPRAPH(document, u'分布的形状参数：'+str(params.shape_para)) # 形状参数
        ADD_PARAPRAPH(document, u'参数单位：'+params.unit) # 参数单位
        ##############################################################################
        document.save(word_name+'.docx')  #保存文档
        #webbrowser.open( word_name+'.docx' )#打开文档
        os.chdir( curr_path ) #返回到当前路径

class SAnalysisWidget(QtWidgets.QWidget, __SAnalysisForm, AbstractFunction):
    '''根据选择模型显示应力应变分析'''
    def __init__(self, master, parent=None):
        super().__init__(parent=parent)

        if not isinstance(master, MainApp):
            raise TypeError
        self.master = master

        self.setupUi(self)
        self.edit_id_to_del.setText('')
        self.edit_id_to_del.setReadOnly(True)

        self.refresh_table()
        self.table_content.clicked.connect(self.chose_one)
        self.btn_submit.clicked.connect(self.show_one)
        self.table_content.contextMenuEvent = self.table_right_click

    def refresh_table(self):

        self.clear_table()

        user = self.master.user
        if user is None:
            return

        s = make_session()
        sim_ls = s.query(SimModel).filter(SimModel.user_id == user.id).all()
        rows = len(sim_ls)
        columns = 4

        self.table_content.setRowCount(rows)
        self.table_content.setColumnCount(columns)

        r = rows - 1
        for sim in sim_ls:

            item = QtWidgets.QTableWidgetItem(str(sim.id))
            item.setTextAlignment(QtCore.Qt.AlignCenter)
            self.table_content.setItem(r, 0, item)

            item = QtWidgets.QTableWidgetItem(str('' if sim.name is None else sim.name))
            item.setTextAlignment(QtCore.Qt.AlignCenter)
            self.table_content.setItem(r, 1, item)

            item = QtWidgets.QTableWidgetItem(str('' if sim.jnl_path is None else sim.jnl_path))
            item.setTextAlignment(QtCore.Qt.AlignCenter)
            self.table_content.setItem(r, 2, item)

            item = QtWidgets.QTableWidgetItem(str('' if sim.external_model_path is None else sim.external_model_path))
            item.setTextAlignment(QtCore.Qt.AlignCenter)
            self.table_content.setItem(r, 3, item)

            r -= 1
        self.edit_id_to_del.setText('')

    def clear_table(self):
        for i in range(self.table_content.rowCount()):
            self.table_content.removeRow(0)

    def chose_current(self):
        row = self.table_content.currentRow()
        item = self.table_content.item(row, 0)
        if not isinstance(item, QtWidgets.QTableWidgetItem):
            return
        self.edit_id_to_del.setText(item.text())

    def chose_one(self):
        self.chose_current()
        
    # 展示图片
    def show_one(self):
        s = make_session()
        _id = self.edit_id_to_del.text().strip()
        try:
            _id = int(_id)
        except (ValueError, TypeError):
            _id = None
        if not _id:
            return self.show_warning_message(message='请先选择一个模型', parent=self)

        sim = s.query(SimModel).filter(SimModel.id == _id).first()
        if sim is None:
            return self.refresh_table()
        s.commit()
        self.refresh_table()

        #获取job name
        fp = os.path.join(base_dir, 'data', 'models', sim.jnl_path)
        from functions.copy_jnl_modify_paras import copy_jnl_modify_paras as __copy_jnl_modify_paras
        content = __copy_jnl_modify_paras(copy_path=fp, ext_path=None, to_add_paras=None)
        job_name = content._job_name()
        #拼接图片路径
        odb_file_path = self.master.temp_path+'\\'+'model_'+str(_id)+'\\'+job_name+'.odb'
        #print(png_file_path)
        if not os.path.exists( odb_file_path ):#ODB文件不存在
            self.show_warning_message(message='仿真结果文件(.odb)不存在，请先仿真！', parent=self)
        else:
            png_file_path = self.master.temp_path+'\\'+'model_'+str(_id)+'\\'+job_name+'_S.png'
            if not os.path.exists( png_file_path ):#应力应变 S 图片文件不存在
                from functions.abaqus_postprocesser import abaqus_postprocesser as __ABAQUS_PostProcesser
                __ABAQUS_PostProcesser(odb_path = odb_file_path).printing_pngs()
            #显示图片
            self.png_show_widget = ShowPNGWidget(png_path = png_file_path)  #
            self.png_show_widget.paintEngine()
            self.png_show_widget.show()
            #self.show_warning_message(message='图片显示', parent=self)
        
    def table_right_click(self, a0):
        context_menu = QtWidgets.QMenu(self)
        show_png_action = QtWidgets.QAction('显示应力应变结果图片', context_menu)

        context_menu.addAction(show_png_action)

        show_png_action.triggered.connect(self.show_one)

        item = self.table_content.currentItem()
        if not isinstance(item, QtWidgets.QTableWidgetItem):
            return
        self.chose_current()

        context_menu.show()
        context_menu.exec_(QtGui.QCursor.pos())


class UAnalysisWidget(QtWidgets.QWidget, __UAnalysisForm, AbstractFunction):
    '''根据选择模型显示应力应变分析'''
    def __init__(self, master, parent=None):
        super().__init__(parent=parent)

        if not isinstance(master, MainApp):
            raise TypeError
        self.master = master

        self.setupUi(self)
        self.edit_id_to_del.setText('')
        self.edit_id_to_del.setReadOnly(True)

        self.refresh_table()
        self.table_content.clicked.connect(self.chose_one)
        self.btn_submit.clicked.connect(self.show_one)
        self.table_content.contextMenuEvent = self.table_right_click

    def refresh_table(self):

        self.clear_table()

        user = self.master.user
        if user is None:
            return

        s = make_session()
        sim_ls = s.query(SimModel).filter(SimModel.user_id == user.id).all()
        rows = len(sim_ls)
        columns = 4

        self.table_content.setRowCount(rows)
        self.table_content.setColumnCount(columns)

        r = rows - 1
        for sim in sim_ls:

            item = QtWidgets.QTableWidgetItem(str(sim.id))
            item.setTextAlignment(QtCore.Qt.AlignCenter)
            self.table_content.setItem(r, 0, item)

            item = QtWidgets.QTableWidgetItem(str('' if sim.name is None else sim.name))
            item.setTextAlignment(QtCore.Qt.AlignCenter)
            self.table_content.setItem(r, 1, item)

            item = QtWidgets.QTableWidgetItem(str('' if sim.jnl_path is None else sim.jnl_path))
            item.setTextAlignment(QtCore.Qt.AlignCenter)
            self.table_content.setItem(r, 2, item)

            item = QtWidgets.QTableWidgetItem(str('' if sim.external_model_path is None else sim.external_model_path))
            item.setTextAlignment(QtCore.Qt.AlignCenter)
            self.table_content.setItem(r, 3, item)

            r -= 1
        self.edit_id_to_del.setText('')

    def clear_table(self):
        for i in range(self.table_content.rowCount()):
            self.table_content.removeRow(0)

    def chose_current(self):
        row = self.table_content.currentRow()
        item = self.table_content.item(row, 0)
        if not isinstance(item, QtWidgets.QTableWidgetItem):
            return
        self.edit_id_to_del.setText(item.text())

    def chose_one(self):
        self.chose_current()
        
    # 展示图片
    def show_one(self):
        s = make_session()
        _id = self.edit_id_to_del.text().strip()
        try:
            _id = int(_id)
        except (ValueError, TypeError):
            _id = None
        if not _id:
            return self.show_warning_message(message='请先选择一个模型', parent=self)

        sim = s.query(SimModel).filter(SimModel.id == _id).first()
        if sim is None:
            return self.refresh_table()
        s.commit()
        self.refresh_table()

        #获取job name
        fp = os.path.join(base_dir, 'data', 'models', sim.jnl_path)
        from functions.copy_jnl_modify_paras import copy_jnl_modify_paras as __copy_jnl_modify_paras
        content = __copy_jnl_modify_paras(copy_path=fp, ext_path=None, to_add_paras=None)
        job_name = content._job_name()
        #拼接图片路径
        odb_file_path = self.master.temp_path+'\\'+'model_'+str(_id)+'\\'+job_name+'.odb'
        #print(png_file_path)
        if not os.path.exists( odb_file_path ):#ODB文件不存在
            self.show_warning_message(message='仿真结果文件(.odb)不存在，请先仿真！', parent=self)
        else:
            png_file_path = self.master.temp_path+'\\'+'model_'+str(_id)+'\\'+job_name+'_U.png'
            if not os.path.exists( png_file_path ):#位移变形 U 图片文件不存在
                from functions.abaqus_postprocesser import abaqus_postprocesser as __ABAQUS_PostProcesser
                __ABAQUS_PostProcesser(odb_path = odb_file_path).printing_pngs()
            #显示图片
            self.png_show_widget = ShowPNGWidget(png_path = png_file_path)  #
            self.png_show_widget.paintEngine()
            self.png_show_widget.show()
            #self.show_warning_message(message='图片显示', parent=self)
        
    def table_right_click(self, a0):
        context_menu = QtWidgets.QMenu(self)
        show_png_action = QtWidgets.QAction('显示位移变形结果图片', context_menu)

        context_menu.addAction(show_png_action)

        show_png_action.triggered.connect(self.show_one)

        item = self.table_content.currentItem()
        if not isinstance(item, QtWidgets.QTableWidgetItem):
            return
        self.chose_current()

        context_menu.show()
        context_menu.exec_(QtGui.QCursor.pos())


from ui.show_png import Ui_Form as __ShowPNGFrom
class ShowPNGWidget(QtWidgets.QWidget, __ShowPNGFrom, AbstractFunction):
    def __init__(self, png_path, parent=None):
        super().__init__(parent=parent)

        self.setupUi(self)
        self.png_path = png_path
        self.setWindowTitle('仿真结果展示')

    def paintEvent(self, event):# set background_img
        painter = QtGui.QPainter(self)
        painter.drawRect(self.rect())
        pixmap = QtGui.QPixmap( self.png_path )
        painter.drawPixmap(self.rect(), pixmap)


class ReliabilityWidget(QtWidgets.QWidget, __ReliabilityForm, AbstractFunction):
    def __init__(self, master, parent=None):
        super().__init__(parent=parent)

        if not isinstance(master, MainApp):
            raise TypeError
        self.master = master

        self.setupUi(self)
        self.edit_id_to_del.setText('')
        self.edit_id_to_del.setReadOnly(True)
        #基于应力应变
        self.btn_mise_reliability.clicked.connect(self.show_mise_reliability)
        self.btn_press_reliability.clicked.connect(self.show_press_reliability)
        self.btn_tresca_reliability.clicked.connect(self.show_tresca_reliability)
        #基于位移变形
        self.btn_x.clicked.connect(self.show_x_reliability)
        self.btn_y.clicked.connect(self.show_y_reliability)
        self.btn_z.clicked.connect(self.show_z_reliability)
        self.btn_mag.clicked.connect(self.show_mag_reliability)

        self.refresh_table()
        self.table_content.clicked.connect(self.chose_one)
        self.table_content.contextMenuEvent = self.table_right_click

    def refresh_table(self):

        self.clear_table()

        user = self.master.user
        if user is None:
            return

        s = make_session()
        sim_ls = s.query(SimModel).filter(SimModel.user_id == user.id).all()
        rows = len(sim_ls)
        columns = 4

        self.table_content.setRowCount(rows)
        self.table_content.setColumnCount(columns)

        r = rows - 1
        for sim in sim_ls:

            item = QtWidgets.QTableWidgetItem(str(sim.id))
            item.setTextAlignment(QtCore.Qt.AlignCenter)
            self.table_content.setItem(r, 0, item)

            item = QtWidgets.QTableWidgetItem(str('' if sim.name is None else sim.name))
            item.setTextAlignment(QtCore.Qt.AlignCenter)
            self.table_content.setItem(r, 1, item)

            item = QtWidgets.QTableWidgetItem(str('' if sim.jnl_path is None else sim.jnl_path))
            item.setTextAlignment(QtCore.Qt.AlignCenter)
            self.table_content.setItem(r, 2, item)

            item = QtWidgets.QTableWidgetItem(str('' if sim.external_model_path is None else sim.external_model_path))
            item.setTextAlignment(QtCore.Qt.AlignCenter)
            self.table_content.setItem(r, 3, item)

            r -= 1
        self.edit_id_to_del.setText('')

    def clear_table(self):
        for i in range(self.table_content.rowCount()):
            self.table_content.removeRow(0)

    def chose_current(self):
        row = self.table_content.currentRow()
        item = self.table_content.item(row, 0)
        if not isinstance(item, QtWidgets.QTableWidgetItem):
            return
        self.edit_id_to_del.setText(item.text())

    def chose_one(self):
        self.chose_current()
        
    def table_right_click(self, a0):
        context_menu = QtWidgets.QMenu(self)
        show_png_action = QtWidgets.QAction('显示位移变形结果图片', context_menu)

        context_menu.addAction(show_png_action)

        show_png_action.triggered.connect(self.show_one)

        item = self.table_content.currentItem()
        if not isinstance(item, QtWidgets.QTableWidgetItem):
            return
        self.chose_current()

        context_menu.show()
        context_menu.exec_(QtGui.QCursor.pos())
    
    def show_mise_reliability(self, ):
        #检查是否选择模型，并传递模型ID
        s = make_session()
        _id = self.edit_id_to_del.text().strip()
        try:
            _id = int(_id)
        except (ValueError, TypeError):
            _id = None
        if not _id:
            return self.show_warning_message(message='请先选择一个模型', parent=self)

        sim = s.query(SimModel).filter(SimModel.id == _id).first()
        if sim is None:
            return self.refresh_table()
        s.commit()
        self.refresh_table()

        #数据地址
        s_txt_path = None
        parent_path = self.master.temp_path+'\\'+'model_'+str(_id)
        if not os.path.exists(parent_path):
            self.show_warning_message("仿真模型路径错误!")
            return 
        filelist = os.listdir( parent_path )
        for filename in filelist:
            file_path = os.path.join(parent_path, filename)
            tmp = filename.split('_')
            if os.path.isfile(file_path) and tmp[-1]=='S.txt': #Specify to find the file
                s_txt_path = file_path # 应力应变数据
        if not s_txt_path:
            self.show_warning_message("数据文件不存在, 请重新仿真!")
            return 
        

        self.mise_widget = MiseFrom(data_path=s_txt_path)
        self.mise_widget.show()
    
    def show_press_reliability(self, ):
        #检查是否选择模型，并传递模型ID
        s = make_session()
        _id = self.edit_id_to_del.text().strip()
        try:
            _id = int(_id)
        except (ValueError, TypeError):
            _id = None
        if not _id:
            return self.show_warning_message(message='请先选择一个模型', parent=self)

        sim = s.query(SimModel).filter(SimModel.id == _id).first()
        if sim is None:
            return self.refresh_table()
        s.commit()
        self.refresh_table()

        #数据地址
        s_txt_path = None
        parent_path = self.master.temp_path+'\\'+'model_'+str(_id)
        if not os.path.exists(parent_path):
            self.show_warning_message("仿真模型路径错误!")
            return 
        filelist = os.listdir( parent_path )
        for filename in filelist:
            file_path = os.path.join(parent_path, filename)
            tmp = filename.split('_')
            if os.path.isfile(file_path) and tmp[-1]=='S.txt': #Specify to find the file
                s_txt_path = file_path # 应力应变数据
        if not s_txt_path:
            self.show_warning_message("数据文件不存在, 请重新仿真!")
            return 

        self.press_widget = PressFrom(data_path=s_txt_path)
        self.press_widget.show()

    def show_tresca_reliability(self, ):
        #检查是否选择模型，并传递模型ID
        s = make_session()
        _id = self.edit_id_to_del.text().strip()
        try:
            _id = int(_id)
        except (ValueError, TypeError):
            _id = None
        if not _id:
            return self.show_warning_message(message='请先选择一个模型', parent=self)

        sim = s.query(SimModel).filter(SimModel.id == _id).first()
        if sim is None:
            return self.refresh_table()
        s.commit()
        self.refresh_table()

        #数据地址
        s_txt_path = None
        parent_path = self.master.temp_path+'\\'+'model_'+str(_id)
        if not os.path.exists(parent_path):
            self.show_warning_message("仿真模型路径错误!")
            return 
        filelist = os.listdir( parent_path )
        for filename in filelist:
            file_path = os.path.join(parent_path, filename)
            tmp = filename.split('_')
            if os.path.isfile(file_path) and tmp[-1]=='S.txt': #Specify to find the file
                s_txt_path = file_path # 应力应变数据
        if not s_txt_path:
            self.show_warning_message("数据文件不存在, 请重新仿真!")
            return 

        self.tresca_widget = TrescaFrom(data_path=s_txt_path)
        self.tresca_widget.show()

    def show_x_reliability(self, ):
        #检查是否选择模型，并传递模型ID
        s = make_session()
        _id = self.edit_id_to_del.text().strip()
        try:
            _id = int(_id)
        except (ValueError, TypeError):
            _id = None
        if not _id:
            return self.show_warning_message(message='请先选择一个模型', parent=self)

        sim = s.query(SimModel).filter(SimModel.id == _id).first()
        if sim is None:
            return self.refresh_table()
        s.commit()
        self.refresh_table()

        #数据地址
        u_txt_path = None
        parent_path = self.master.temp_path+'\\'+'model_'+str(_id)
        if not os.path.exists(parent_path):
            self.show_warning_message("仿真模型路径错误!")
            return 
        filelist = os.listdir( parent_path )
        for filename in filelist:
            file_path = os.path.join(parent_path, filename)
            tmp = filename.split('_')
            if os.path.isfile(file_path) and tmp[-1]=='U.txt': #Specify to find the file
                u_txt_path = file_path # 应力应变数据
        if not u_txt_path:
            self.show_warning_message("数据文件不存在, 请重新仿真!")
            return 

        self.x_widget = XFrom(data_path=u_txt_path)
        self.x_widget.show()

    def show_y_reliability(self, ):
        #检查是否选择模型，并传递模型ID
        s = make_session()
        _id = self.edit_id_to_del.text().strip()
        try:
            _id = int(_id)
        except (ValueError, TypeError):
            _id = None
        if not _id:
            return self.show_warning_message(message='请先选择一个模型', parent=self)

        sim = s.query(SimModel).filter(SimModel.id == _id).first()
        if sim is None:
            return self.refresh_table()
        s.commit()
        self.refresh_table()

        #数据地址
        u_txt_path = None
        parent_path = self.master.temp_path+'\\'+'model_'+str(_id)
        if not os.path.exists(parent_path):
            self.show_warning_message("仿真模型路径错误!")
            return 
        filelist = os.listdir( parent_path )
        for filename in filelist:
            file_path = os.path.join(parent_path, filename)
            tmp = filename.split('_')
            if os.path.isfile(file_path) and tmp[-1]=='U.txt': #Specify to find the file
                u_txt_path = file_path # 应力应变数据
        if not u_txt_path:
            self.show_warning_message("数据文件不存在, 请重新仿真!")
            return 

        self.y_widget = YFrom(data_path=u_txt_path)
        self.y_widget.show()

    def show_z_reliability(self, ):
        #检查是否选择模型，并传递模型ID
        s = make_session()
        _id = self.edit_id_to_del.text().strip()
        try:
            _id = int(_id)
        except (ValueError, TypeError):
            _id = None
        if not _id:
            return self.show_warning_message(message='请先选择一个模型', parent=self)

        sim = s.query(SimModel).filter(SimModel.id == _id).first()
        if sim is None:
            return self.refresh_table()
        s.commit()
        self.refresh_table()

        #数据地址
        u_txt_path = None
        parent_path = self.master.temp_path+'\\'+'model_'+str(_id)
        if not os.path.exists(parent_path):
            self.show_warning_message("仿真模型路径错误!")
            return 
        filelist = os.listdir( parent_path )
        for filename in filelist:
            file_path = os.path.join(parent_path, filename)
            tmp = filename.split('_')
            if os.path.isfile(file_path) and tmp[-1]=='U.txt': #Specify to find the file
                u_txt_path = file_path # 应力应变数据
        if not u_txt_path:
            self.show_warning_message("数据文件不存在, 请重新仿真!")
            return 

        self.z_widget = ZFrom(data_path=u_txt_path)
        self.z_widget.show()

    def show_mag_reliability(self, ):
        #检查是否选择模型，并传递模型ID
        s = make_session()
        _id = self.edit_id_to_del.text().strip()
        try:
            _id = int(_id)
        except (ValueError, TypeError):
            _id = None
        if not _id:
            return self.show_warning_message(message='请先选择一个模型', parent=self)

        sim = s.query(SimModel).filter(SimModel.id == _id).first()
        if sim is None:
            return self.refresh_table()
        s.commit()
        self.refresh_table()

        #数据地址
        u_txt_path = None
        parent_path = self.master.temp_path+'\\'+'model_'+str(_id)
        if not os.path.exists(parent_path):
            self.show_warning_message("仿真模型路径错误!")
            return 
        filelist = os.listdir( parent_path )
        for filename in filelist:
            file_path = os.path.join(parent_path, filename)
            tmp = filename.split('_')
            if os.path.isfile(file_path) and tmp[-1]=='U.txt': #Specify to find the file
                u_txt_path = file_path # 应力应变数据
        if not u_txt_path:
            self.show_warning_message("数据文件不存在, 请重新仿真!")
            return 

        self.mag_widget = MagFrom(data_path=u_txt_path)
        self.mag_widget.show()

from ui.reliability_mise_form import Ui_Form as __MiseFrom
from ui.reliability_press_form import Ui_Form as __PressFrom
from ui.reliability_tresca_form import Ui_Form as __TrescaFrom
from ui.reliability_x_form import Ui_Form as __XFrom
from ui.reliability_y_form import Ui_Form as __YFrom
from ui.reliability_z_form import Ui_Form as __ZFrom
from ui.reliability_u_mag_form import Ui_Form as __MagFrom

#coding:utf-8
import matplotlib.pyplot as plt
plt.rcParams['font.sans-serif']=['SimHei'] #用来正常显示中文标签
plt.rcParams['axes.unicode_minus']=False #用来正常显示负号
#有中文出现的情况，需要u'内容'

class MiseFrom(QtWidgets.QWidget, __MiseFrom, AbstractFunction):
    def __init__(self, data_path, parent=None):
        super().__init__(parent=parent)
        self.setupUi(self)
        self.setWindowTitle('基于应力的可靠性分析')

        self.pushButton_3.clicked.connect(self.computing_threshold_reliability)
        self.data_path = data_path

        import pandas as pd
        try:
            data = pd.read_csv(self.data_path, sep=',') #.txt,分隔符为','
            self.data = data[' mises']
        except:
            self.show_warning_message(message="数据读取失败", title='警告', parent=self)

        self.show_data_pics()
        
    def computing_threshold_reliability(self):
        threshold = self.value_input.value()
        r = len(self.data[ self.data>threshold ])/len(self.data)
        message = "阈值为"+format(threshold, '.5f')+"时，可靠度为"+format(r, '.2%')
        self.show_warning_message(message=message, title='阈值可靠度', parent=self)

    def show_data_pics(self):
        parent_path = self.data_path.split('\\')
        parent_path.pop(-1)
        parent_path = '\\'.join( parent_path )
        import seaborn as sns
        if not os.path.isfile(parent_path+"\\_mise_hist.png"):
            hist = sns.distplot(self.data)
            hist.set_xlabel(u'应力值(牛顿/N)')
            hist.set_ylabel(u'单元个数')
            hist.spines['top'].set_visible(False) #去掉上边框
            hist.spines['right'].set_visible(False) #去掉右边框
            hist_fig = hist.get_figure()
            hist_fig.savefig(parent_path+"\\_mise_hist.png")
            hist_fig.clear()
        if not os.path.isfile(parent_path+"\\_mise_threshold_relia.png"):
            threshold_relia = sns.kdeplot(self.data, cumulative=True, legend=False)
            threshold_relia.set_xlabel(u'应力值(牛顿/N)')
            threshold_relia.set_ylabel(u'可靠度')
            #threshold_relia.set_ylim(bottom=0.0, top=1.0)
            threshold_relia.invert_xaxis()
            #threshold_relia.set_xbound(lower=self.data.min(), upper=self.data.max())
            threshold_relia.spines['top'].set_visible(False) #去掉上边框
            threshold_relia.spines['right'].set_visible(False) #去掉右边框
            threshold_relia.ticklabel_format() #横纵坐标科学计数
            threshold_relia_fig = threshold_relia.get_figure()
            threshold_relia_fig.savefig(parent_path+"\\_mise_threshold_relia.png")
            #翻转刻度标记
            x_tick_labels = []
            for _ in threshold_relia.get_xticklabels():
                s = _.get_text()
                x_tick_labels.append( s )
            threshold_relia.set_xticklabels( x_tick_labels[::-1] )
            threshold_relia_fig = threshold_relia.get_figure()
            threshold_relia_fig.savefig(parent_path+"\\_mise_threshold_relia.png")
            threshold_relia_fig.clear()

        hist_pix = QtGui.QPixmap(parent_path+"\\_mise_hist.png")
        threshold_relia_pix = QtGui.QPixmap(parent_path+"\\_mise_threshold_relia.png")
        self.pdf_pic.setPixmap( hist_pix )
        self.pdf_pic.setScaledContents(True)
        self.reliability_pic.setPixmap( threshold_relia_pix )
        self.reliability_pic.setScaledContents(True)
        

class PressFrom(QtWidgets.QWidget, __PressFrom, AbstractFunction):
    def __init__(self, data_path, parent=None):
        super().__init__(parent=parent)
        self.setupUi(self)
        self.setWindowTitle('基于压力的可靠性分析')

        self.pushButton_3.clicked.connect(self.computing_threshold_reliability)
        self.data_path = data_path

        import pandas as pd
        try:
            data = pd.read_csv(data_path, sep=',') #.txt,分隔符为','
            self.data = data[' press']
        except:
            self.show_warning_message(message="数据读取失败", title='警告', parent=self)

        self.show_data_pics()
        
    def computing_threshold_reliability(self):
        threshold = self.value_input.value()
        r = len(self.data[ self.data>threshold ])/len(self.data)
        message = "阈值为"+format(threshold, '.5f')+"时，可靠度为"+format(r, '.2%')
        self.show_warning_message(message=message, title='阈值可靠度', parent=self)

    def show_data_pics(self):
        parent_path = self.data_path.split('\\')
        parent_path.pop(-1)
        parent_path = '\\'.join( parent_path )
        import seaborn as sns
        if not os.path.isfile(parent_path+"\\_press_hist.png"):
            hist = sns.distplot(self.data)
            hist.set_xlabel(u'压力值(帕/Pa)')
            hist.set_ylabel(u'单元个数')
            hist.spines['top'].set_visible(False) #去掉上边框
            hist.spines['right'].set_visible(False) #去掉右边框
            hist_fig = hist.get_figure()
            hist_fig.savefig(parent_path+"\\_press_hist.png")
            hist_fig.clear()
        if not os.path.isfile(parent_path+"\\_press_threshold_relia.png"):
            threshold_relia = sns.kdeplot(self.data, cumulative=True, legend=False)
            threshold_relia.set_xlabel(u'压力值(帕/Pa)')
            threshold_relia.set_ylabel(u'可靠度')
            #threshold_relia.set_ylim(bottom=0.0, top=1.0)
            threshold_relia.invert_xaxis()
            #threshold_relia.set_xbound(lower=self.data.min(), upper=self.data.max())
            threshold_relia.spines['top'].set_visible(False) #去掉上边框
            threshold_relia.spines['right'].set_visible(False) #去掉右边框
            threshold_relia.ticklabel_format() #横纵坐标科学计数
            threshold_relia_fig = threshold_relia.get_figure()
            threshold_relia_fig.savefig(parent_path+"\\_press_threshold_relia.png")
            #翻转刻度标记
            x_tick_labels = []
            for _ in threshold_relia.get_xticklabels():
                s = _.get_text()
                x_tick_labels.append( s )
            threshold_relia.set_xticklabels( x_tick_labels[::-1] )
            threshold_relia_fig = threshold_relia.get_figure()
            threshold_relia_fig.savefig(parent_path+"\\_press_threshold_relia.png")
            threshold_relia_fig.clear()

        hist_pix = QtGui.QPixmap(parent_path+"\\_press_hist.png")
        threshold_relia_pix = QtGui.QPixmap(parent_path+"\\_press_threshold_relia.png")
        self.pdf_pic.setPixmap( hist_pix )
        self.pdf_pic.setScaledContents(True)
        self.reliability_pic.setPixmap( threshold_relia_pix )
        self.reliability_pic.setScaledContents(True)

class TrescaFrom(QtWidgets.QWidget, __TrescaFrom, AbstractFunction):
    def __init__(self, data_path, parent=None):
        super().__init__(parent=parent)
        self.setupUi(self)
        self.setWindowTitle('基于屈服的可靠性分析')

        self.pushButton_3.clicked.connect(self.computing_threshold_reliability)
        self.data_path = data_path

        import pandas as pd
        try:
            data = pd.read_csv(data_path, sep=',') #.txt,分隔符为','
            self.data = data[' tresca']
        except:
            self.show_warning_message(message="数据读取失败", title='警告', parent=self)

        self.show_data_pics()
        
    def computing_threshold_reliability(self):
        threshold = self.value_input.value()
        r = len(self.data[ self.data>threshold ])/len(self.data)
        message = "阈值为"+format(threshold, '.5f')+"时，可靠度为"+format(r, '.2%')
        self.show_warning_message(message=message, title='阈值可靠度', parent=self)

    def show_data_pics(self):
        parent_path = self.data_path.split('\\')
        parent_path.pop(-1)
        parent_path = '\\'.join( parent_path )
        import seaborn as sns
        if not os.path.isfile(parent_path+"\\_tresca_hist.png"):
            hist = sns.distplot(self.data)
            hist.set_xlabel(u'屈服值(帕/Pa)')
            hist.set_ylabel(u'单元个数')
            hist.spines['top'].set_visible(False) #去掉上边框
            hist.spines['right'].set_visible(False) #去掉右边框
            hist_fig = hist.get_figure()
            hist_fig.savefig(parent_path+"\\_tresca_hist.png")
            hist_fig.clear()
        if not os.path.isfile(parent_path+"\\_tresca_threshold_relia.png"):
            threshold_relia = sns.kdeplot(self.data, cumulative=True, legend=False)
            threshold_relia.set_xlabel(u'屈服值(帕/Pa)')
            threshold_relia.set_ylabel(u'可靠度')
            #threshold_relia.set_ylim(bottom=0.0, top=1.0)
            threshold_relia.invert_xaxis()
            threshold_relia.set_xbound(lower=self.data.min(), upper=self.data.max())
            threshold_relia.spines['top'].set_visible(False) #去掉上边框
            threshold_relia.spines['right'].set_visible(False) #去掉右边框
            threshold_relia.ticklabel_format() #横纵坐标科学计数
            threshold_relia_fig = threshold_relia.get_figure()
            threshold_relia_fig.savefig(parent_path+"\\_tresca_threshold_relia.png")
            #翻转刻度标记
            x_tick_labels = []
            for _ in threshold_relia.get_xticklabels():
                s = _.get_text()
                x_tick_labels.append( s )
            threshold_relia.set_xticklabels( x_tick_labels[::-1] )
            threshold_relia_fig = threshold_relia.get_figure()
            threshold_relia_fig.savefig(parent_path+"\\_tresca_threshold_relia.png")
            threshold_relia_fig.clear()

        hist_pix = QtGui.QPixmap(parent_path+"\\_tresca_hist.png")
        threshold_relia_pix = QtGui.QPixmap(parent_path+"\\_tresca_threshold_relia.png")
        self.pdf_pic.setPixmap( hist_pix )
        self.pdf_pic.setScaledContents(True)
        self.reliability_pic.setPixmap( threshold_relia_pix )
        self.reliability_pic.setScaledContents(True)

class XFrom(QtWidgets.QWidget, __XFrom, AbstractFunction):
    def __init__(self, data_path, parent=None):
        super().__init__(parent=parent)
        self.setupUi(self)
        self.setWindowTitle('基于X方向位移的可靠性分析')

        self.pushButton_3.clicked.connect(self.computing_threshold_reliability)
        self.data_path = data_path

        import pandas as pd
        try:
            data = pd.read_csv(data_path, sep=',') #.txt,分隔符为','
            self.data = abs(data[' X'])
        except:
            self.show_warning_message(message="数据读取失败", title='警告', parent=self)

        self.show_data_pics()
        
    def computing_threshold_reliability(self):
        threshold = self.value_input.value()
        r = len(self.data[ self.data>threshold ])/len(self.data)
        message = "阈值为"+format(threshold, '.10f')+"时，可靠度为"+format(r, '.2%')
        self.show_warning_message(message=message, title='阈值可靠度', parent=self)

    def show_data_pics(self):
        parent_path = self.data_path.split('\\')
        parent_path.pop(-1)
        parent_path = '\\'.join( parent_path )
        import seaborn as sns
        if not os.path.isfile(parent_path+"\\_x_hist.png"):
            hist = sns.distplot(self.data)
            hist.set_xlabel(u'X方向位移量(米/m)')
            hist.set_ylabel(u'节点个数')
            hist.spines['top'].set_visible(False) #去掉上边框
            hist.spines['right'].set_visible(False) #去掉右边框
            hist_fig = hist.get_figure()
            hist_fig.savefig(parent_path+"\\_x_hist.png")
            hist_fig.clear()
        if not os.path.isfile(parent_path+"\\_x_threshold_relia.png"):
            threshold_relia = sns.kdeplot(self.data, cumulative=True, legend=False)
            threshold_relia.set_xlabel(u'X方向位移量(米/m)')
            threshold_relia.set_ylabel(u'可靠度')
            #threshold_relia.set_ylim(bottom=0.0, top=1.0)
            threshold_relia.invert_xaxis()
            #threshold_relia.set_xbound(lower=self.data.min(), upper=self.data.max())
            threshold_relia.spines['top'].set_visible(False) #去掉上边框
            threshold_relia.spines['right'].set_visible(False) #去掉右边框
            threshold_relia.ticklabel_format() #横纵坐标科学计数
            threshold_relia_fig = threshold_relia.get_figure()
            threshold_relia_fig.savefig(parent_path+"\\_x_threshold_relia.png")
            #翻转刻度标记
            x_tick_labels = []
            for _ in threshold_relia.get_xticklabels():
                s = _.get_text()
                x_tick_labels.append( s )
            threshold_relia.set_xticklabels( x_tick_labels[::-1] )
            threshold_relia_fig = threshold_relia.get_figure()
            threshold_relia_fig.savefig(parent_path+"\\_x_threshold_relia.png")
            threshold_relia_fig.clear()

        hist_pix = QtGui.QPixmap(parent_path+"\\_x_hist.png")
        threshold_relia_pix = QtGui.QPixmap(parent_path+"\\_x_threshold_relia.png")
        self.pdf_pic.setPixmap( hist_pix )
        self.pdf_pic.setScaledContents(True)
        self.reliability_pic.setPixmap( threshold_relia_pix )
        self.reliability_pic.setScaledContents(True)

class YFrom(QtWidgets.QWidget, __YFrom, AbstractFunction):
    def __init__(self, data_path, parent=None):
        super().__init__(parent=parent)
        self.setupUi(self)
        self.setWindowTitle('基于Y方向位移的可靠性分析')

        self.pushButton_3.clicked.connect(self.computing_threshold_reliability)
        self.data_path = data_path

        import pandas as pd
        try:
            data = pd.read_csv(data_path, sep=',') #.txt,分隔符为','
            self.data = abs(data[' Y'])
        except:
            self.show_warning_message(message="数据读取失败", title='警告', parent=self)

        self.show_data_pics()
        
    def computing_threshold_reliability(self):
        threshold = self.value_input.value()
        r = len(self.data[ self.data>threshold ])/len(self.data)
        message = "阈值为"+format(threshold, '.10f')+"时，可靠度为"+format(r, '.2%')
        self.show_warning_message(message=message, title='阈值可靠度', parent=self)

    def show_data_pics(self):
        parent_path = self.data_path.split('\\')
        parent_path.pop(-1)
        parent_path = '\\'.join( parent_path )
        import seaborn as sns
        if not os.path.isfile(parent_path+"\\_y_hist.png"):
            hist = sns.distplot(self.data)
            hist.set_xlabel(u'Y方向位移量(米/m)')
            hist.set_ylabel(u'节点个数')
            hist.spines['top'].set_visible(False) #去掉上边框
            hist.spines['right'].set_visible(False) #去掉右边框
            hist_fig = hist.get_figure()
            hist_fig.savefig(parent_path+"\\_y_hist.png")
            hist_fig.clear()
        if not os.path.isfile(parent_path+"\\_y_threshold_relia.png"):
            threshold_relia = sns.kdeplot(self.data, cumulative=True, legend=False)
            threshold_relia.set_xlabel(u'Y方向位移量(米/m)')
            threshold_relia.set_ylabel(u'可靠度')
            #threshold_relia.set_ylim(bottom=0.0, top=1.0)
            threshold_relia.invert_xaxis()
            #threshold_relia.set_xbound(lower=self.data.min(), upper=self.data.max())
            threshold_relia.spines['top'].set_visible(False) #去掉上边框
            threshold_relia.spines['right'].set_visible(False) #去掉右边框
            threshold_relia.ticklabel_format() #横纵坐标科学计数
            threshold_relia_fig = threshold_relia.get_figure()
            threshold_relia_fig.savefig(parent_path+"\\_y_threshold_relia.png")
            #翻转刻度标记
            x_tick_labels = []
            for _ in threshold_relia.get_xticklabels():
                s = _.get_text()
                x_tick_labels.append( s )
            threshold_relia.set_xticklabels( x_tick_labels[::-1] )
            threshold_relia_fig = threshold_relia.get_figure()
            threshold_relia_fig.savefig(parent_path+"\\_y_threshold_relia.png")
            threshold_relia_fig.clear()

        hist_pix = QtGui.QPixmap(parent_path+"\\_y_hist.png")
        threshold_relia_pix = QtGui.QPixmap(parent_path+"\\_y_threshold_relia.png")
        self.pdf_pic.setPixmap( hist_pix )
        self.pdf_pic.setScaledContents(True)
        self.reliability_pic.setPixmap( threshold_relia_pix )
        self.reliability_pic.setScaledContents(True)

class ZFrom(QtWidgets.QWidget, __ZFrom, AbstractFunction):
    def __init__(self, data_path, parent=None):
        super().__init__(parent=parent)
        self.setupUi(self)
        self.setWindowTitle('基于Z方向位移的可靠性分析')

        self.pushButton_3.clicked.connect(self.computing_threshold_reliability)
        self.data_path = data_path

        import pandas as pd
        try:
            data = pd.read_csv(data_path, sep=',') #.txt,分隔符为','
            self.data = abs(data[' Z'])
        except:
            self.show_warning_message(message="数据读取失败", title='警告', parent=self)

        self.show_data_pics()
        
    def computing_threshold_reliability(self):
        threshold = self.value_input.value()
        r = len(self.data[ self.data>threshold ])/len(self.data)
        message = "阈值为"+format(threshold, '.10f')+"时，可靠度为"+format(r, '.2%')
        self.show_warning_message(message=message, title='阈值可靠度', parent=self)

    def show_data_pics(self):
        parent_path = self.data_path.split('\\')
        parent_path.pop(-1)
        parent_path = '\\'.join( parent_path )
        import seaborn as sns
        if not os.path.isfile(parent_path+"\\_z_hist.png"):
            hist = sns.distplot(self.data)
            hist.set_xlabel(u'Z方向位移量(米/m)')
            hist.set_ylabel(u'节点个数')
            hist.spines['top'].set_visible(False) #去掉上边框
            hist.spines['right'].set_visible(False) #去掉右边框
            hist_fig = hist.get_figure()
            hist_fig.savefig(parent_path+"\\_z_hist.png")
            hist_fig.clear()
        if not os.path.isfile(parent_path+"\\_z_threshold_relia.png"):
            threshold_relia = sns.kdeplot(self.data, cumulative=True, legend=False)
            threshold_relia.set_xlabel(u'Z方向位移量(米/m)')
            threshold_relia.set_ylabel(u'可靠度')
            #threshold_relia.set_ylim(bottom=0.0, top=1.0)
            threshold_relia.invert_xaxis()
            #threshold_relia.set_xbound(lower=self.data.min(), upper=self.data.max())
            threshold_relia.spines['top'].set_visible(False) #去掉上边框
            threshold_relia.spines['right'].set_visible(False) #去掉右边框
            threshold_relia.ticklabel_format() #横纵坐标科学计数
            threshold_relia_fig = threshold_relia.get_figure()
            threshold_relia_fig.savefig(parent_path+"\\_z_threshold_relia.png")
            #翻转刻度标记
            x_tick_labels = []
            for _ in threshold_relia.get_xticklabels():
                s = _.get_text()
                x_tick_labels.append( s )
            threshold_relia.set_xticklabels( x_tick_labels[::-1] )
            threshold_relia_fig = threshold_relia.get_figure()
            threshold_relia_fig.savefig(parent_path+"\\_z_threshold_relia.png")
            threshold_relia_fig.clear()

        hist_pix = QtGui.QPixmap(parent_path+"\\_z_hist.png")
        threshold_relia_pix = QtGui.QPixmap(parent_path+"\\_z_threshold_relia.png")
        self.pdf_pic.setPixmap( hist_pix )
        self.pdf_pic.setScaledContents(True)
        self.reliability_pic.setPixmap( threshold_relia_pix )
        self.reliability_pic.setScaledContents(True)

class MagFrom(QtWidgets.QWidget, __MagFrom, AbstractFunction):
    def __init__(self, data_path, parent=None):
        super().__init__(parent=parent)
        self.setupUi(self)
        self.setWindowTitle('基于总位移量的可靠性分析')

        self.pushButton_3.clicked.connect(self.computing_threshold_reliability)
        self.data_path = data_path

        import pandas as pd
        try:
            data = pd.read_csv(data_path, sep=',') #.txt,分隔符为','
            self.data = abs(data[' magnitude'])
        except:
            self.show_warning_message(message="数据读取失败", title='警告', parent=self)

        self.show_data_pics()
        
    def computing_threshold_reliability(self):
        threshold = self.value_input.value()
        r = len(self.data[ self.data>threshold ])/len(self.data)
        message = "阈值为"+format(threshold, '.10f')+"时，可靠度为"+format(r, '.2%')
        self.show_warning_message(message=message, title='阈值可靠度', parent=self)

    def show_data_pics(self):
        parent_path = self.data_path.split('\\')
        parent_path.pop(-1)
        parent_path = '\\'.join( parent_path )
        import seaborn as sns
        if not os.path.isfile(parent_path+"\\_mag_hist.png"):
            hist = sns.distplot(self.data)
            hist.set_xlabel(u'总位移量(米/m)')
            hist.set_ylabel(u'节点个数')
            hist.spines['top'].set_visible(False) #去掉上边框
            hist.spines['right'].set_visible(False) #去掉右边框
            hist_fig = hist.get_figure()
            hist_fig.savefig(parent_path+"\\_mag_hist.png")
            hist_fig.clear()
        if not os.path.isfile(parent_path+"\\_mag_threshold_relia.png"):
            threshold_relia = sns.kdeplot(self.data, cumulative=True, legend=False)
            threshold_relia.set_xlabel(u'总位移量(米/m)')
            threshold_relia.set_ylabel(u'可靠度')
            #threshold_relia.set_ylim(bottom=0.0, top=1.0)
            threshold_relia.invert_xaxis()
            #threshold_relia.set_xbound(lower=self.data.min(), upper=self.data.max())
            threshold_relia.spines['top'].set_visible(False) #去掉上边框
            threshold_relia.spines['right'].set_visible(False) #去掉右边框
            threshold_relia.ticklabel_format() #横纵坐标科学计数
            threshold_relia_fig = threshold_relia.get_figure()
            threshold_relia_fig.savefig(parent_path+"\\_mag_threshold_relia.png")
            #翻转刻度标记
            x_tick_labels = []
            for _ in threshold_relia.get_xticklabels():
                s = _.get_text()
                x_tick_labels.append( s )
            threshold_relia.set_xticklabels( x_tick_labels[::-1] )
            threshold_relia_fig = threshold_relia.get_figure()
            threshold_relia_fig.savefig(parent_path+"\\_mag_threshold_relia.png")
            threshold_relia_fig.clear()

        hist_pix = QtGui.QPixmap(parent_path+"\\_mag_hist.png")
        threshold_relia_pix = QtGui.QPixmap(parent_path+"\\_mag_threshold_relia.png")
        self.pdf_pic.setPixmap( hist_pix )
        self.pdf_pic.setScaledContents(True)
        self.reliability_pic.setPixmap( threshold_relia_pix )
        self.reliability_pic.setScaledContents(True)



class ReportWidget(QtWidgets.QWidget, __ReportForm, AbstractFunction):
    '''生成报告并显示'''
    def __init__(self, master, parent=None):
        super().__init__(parent=parent)

        if not isinstance(master, MainApp):
            raise TypeError
        self.master = master

        self.setupUi(self)
        self.edit_id_to_del.setText('')
        self.edit_id_to_del.setReadOnly(True)

        self.refresh_table()
        self.table_content.clicked.connect(self.chose_one)
        self.btn_submit.clicked.connect(self.show_one)
        self.table_content.contextMenuEvent = self.table_right_click

    def refresh_table(self):

        self.clear_table()

        user = self.master.user
        if user is None:
            return

        s = make_session()
        sim_ls = s.query(SimModel).filter(SimModel.user_id == user.id).all()
        rows = len(sim_ls)
        columns = 4

        self.table_content.setRowCount(rows)
        self.table_content.setColumnCount(columns)

        r = rows - 1
        for sim in sim_ls:

            item = QtWidgets.QTableWidgetItem(str(sim.id))
            item.setTextAlignment(QtCore.Qt.AlignCenter)
            self.table_content.setItem(r, 0, item)

            item = QtWidgets.QTableWidgetItem(str('' if sim.name is None else sim.name))
            item.setTextAlignment(QtCore.Qt.AlignCenter)
            self.table_content.setItem(r, 1, item)

            item = QtWidgets.QTableWidgetItem(str('' if sim.jnl_path is None else sim.jnl_path))
            item.setTextAlignment(QtCore.Qt.AlignCenter)
            self.table_content.setItem(r, 2, item)

            item = QtWidgets.QTableWidgetItem(str('' if sim.external_model_path is None else sim.external_model_path))
            item.setTextAlignment(QtCore.Qt.AlignCenter)
            self.table_content.setItem(r, 3, item)

            r -= 1
        self.edit_id_to_del.setText('')

    def clear_table(self):
        for i in range(self.table_content.rowCount()):
            self.table_content.removeRow(0)

    def chose_current(self):
        row = self.table_content.currentRow()
        item = self.table_content.item(row, 0)
        if not isinstance(item, QtWidgets.QTableWidgetItem):
            return
        self.edit_id_to_del.setText(item.text())

    def chose_one(self):
        self.chose_current()
        
    def table_right_click(self, a0):
        context_menu = QtWidgets.QMenu(self)
        show_png_action = QtWidgets.QAction('显示位移变形结果图片', context_menu)

        context_menu.addAction(show_png_action)

        show_png_action.triggered.connect(self.show_one)

        item = self.table_content.currentItem()
        if not isinstance(item, QtWidgets.QTableWidgetItem):
            return
        self.chose_current()

        context_menu.show()
        context_menu.exec_(QtGui.QCursor.pos())
    
        # 展示报告
    
    def show_one(self):
        s = make_session()
        _id = self.edit_id_to_del.text().strip()
        try:
            _id = int(_id)
        except (ValueError, TypeError):
            _id = None
        if not _id:
            return self.show_warning_message(message='请先选择一个模型', parent=self)

        sim = s.query(SimModel).filter(SimModel.id == _id).first()
        if sim is None:
            return self.refresh_table()
        s.commit()
        self.refresh_table()

        #获取job name
        fp = os.path.join(base_dir, 'data', 'models', sim.jnl_path)
        from functions.copy_jnl_modify_paras import copy_jnl_modify_paras as __copy_jnl_modify_paras
        content = __copy_jnl_modify_paras(copy_path=fp, ext_path=None, to_add_paras=None)
        job_name = content._job_name()
        #云图图片路径
        odb_file_path = self.master.temp_path+'\\'+'model_'+str(_id)+'\\'+job_name+'.odb'
        u_png_file_path = self.master.temp_path+'\\'+'model_'+str(_id)+'\\'+job_name+'_U.png' # 位移变形云图
        s_png_file_path = self.master.temp_path+'\\'+'model_'+str(_id)+'\\'+job_name+'_S.png' # 应力应变云图
        #print(png_file_path)
        if not os.path.exists( odb_file_path ):#ODB文件不存在
            self.show_warning_message(message='仿真结果文件(.odb)不存在，请先仿真！', parent=self)
            return 
        if not os.path.exists( u_png_file_path ) or not os.path.exists( s_png_file_path ): #云图不存在
            self.show_warning_message(message='位移变形(或应力应变)云图错误', parent=self)
            return 

        # mise
        pdf_mise_png_path = self.master.temp_path+'\\'+'model_'+str(_id)+'\\'+'_mise'+'_hist.png'
        relia_mise_png_path = self.master.temp_path+'\\'+'model_'+str(_id)+'\\'+'_mise'+'_threshold_relia.png'
        if not os.path.exists( pdf_mise_png_path ) or \
            not os.path.exists( relia_mise_png_path ):
            self.show_warning_message(message='应力值可靠性分析缺失', parent=self)
            return 
        # press
        pdf_press_png_path = self.master.temp_path+'\\'+'model_'+str(_id)+'\\'+'_press'+'_hist.png'
        relia_press_png_path = self.master.temp_path+'\\'+'model_'+str(_id)+'\\'+'_press'+'_threshold_relia.png'
        if not os.path.exists( pdf_press_png_path ) or \
            not os.path.exists( relia_press_png_path ):
            self.show_warning_message(message='压力值可靠性分析缺失', parent=self)
            return 
        # tresca
        pdf_tresca_png_path = self.master.temp_path+'\\'+'model_'+str(_id)+'\\'+'_tresca'+'_hist.png'
        relia_tresca_png_path = self.master.temp_path+'\\'+'model_'+str(_id)+'\\'+'_tresca'+'_threshold_relia.png'
        if not os.path.exists( pdf_tresca_png_path ) or \
            not os.path.exists( relia_tresca_png_path ):
            self.show_warning_message(message='屈服值可靠性分析缺失', parent=self)
            return 
        # x
        pdf_x_png_path = self.master.temp_path+'\\'+'model_'+str(_id)+'\\'+'_x'+'_hist.png'
        relia_x_png_path = self.master.temp_path+'\\'+'model_'+str(_id)+'\\'+'_x'+'_threshold_relia.png'
        if not os.path.exists( pdf_x_png_path ) or \
            not os.path.exists( relia_x_png_path ):
            self.show_warning_message(message='X轴位移量可靠性分析缺失', parent=self)
            return 
        # y
        pdf_y_png_path = self.master.temp_path+'\\'+'model_'+str(_id)+'\\'+'_y'+'_hist.png'
        relia_y_png_path = self.master.temp_path+'\\'+'model_'+str(_id)+'\\'+'_y'+'_threshold_relia.png'
        if not os.path.exists( pdf_y_png_path ) or \
            not os.path.exists( relia_y_png_path ):
            self.show_warning_message(message='Y轴位移量可靠性分析缺失', parent=self)
            return 
        # z
        pdf_z_png_path = self.master.temp_path+'\\'+'model_'+str(_id)+'\\'+'_z'+'_hist.png'
        relia_z_png_path = self.master.temp_path+'\\'+'model_'+str(_id)+'\\'+'_z'+'_threshold_relia.png'
        if not os.path.exists( pdf_z_png_path ) or \
            not os.path.exists( relia_z_png_path ):
            self.show_warning_message(message='Z轴位移量可靠性分析缺失', parent=self)
            return 
        # mag
        pdf_mag_png_path = self.master.temp_path+'\\'+'model_'+str(_id)+'\\'+'_mag'+'_hist.png'
        relia_mag_png_path = self.master.temp_path+'\\'+'model_'+str(_id)+'\\'+'_mag'+'_threshold_relia.png'
        if not os.path.exists( pdf_mag_png_path ) or \
            not os.path.exists( relia_mag_png_path ):
            self.show_warning_message(message='总位移量可靠性分析缺失', parent=self)
            return 
            
        # 输出数据到txt
        u_txt_path = self.master.temp_path+'\\'+'model_'+str(_id)+'\\'+job_name+'_U.txt' # 位移变形数据
        s_txt_path = self.master.temp_path+'\\'+'model_'+str(_id)+'\\'+job_name+'_S.txt' # 应力应变数据
        if not os.path.exists( u_txt_path ) or not os.path.exists( s_txt_path ):
            from functions.abaqus_postprocesser import abaqus_postprocesser as __ABAQUS_PostProcesser
            __ABAQUS_PostProcesser(odb_path = odb_file_path).reading_data()
        # 图片
        pngs_path_dict = {  'u_cloud_png':u_png_file_path, 's_cloud_png':s_png_file_path, # 云图
                            'pdf_mise':pdf_mise_png_path, 'pdf_press':pdf_press_png_path, #PDF图
                            'pdf_tresca':pdf_tresca_png_path, 'pdf_x':pdf_x_png_path, 
                            'pdf_y':pdf_y_png_path, 'pdf_z':pdf_z_png_path, 
                            'pdf_mag':pdf_mag_png_path, 
                            'relia_mise':relia_mise_png_path, 'relia_press':relia_press_png_path, #可靠性图
                            'relia_tresca':relia_tresca_png_path, 'relia_x':relia_x_png_path, 
                            'relia_y':relia_y_png_path, 'relia_z':relia_z_png_path, 
                            'relia_mag':relia_mag_png_path, }
        # 数据分析
        data_analysis_dict = { 'u_analysis':self.data_analysis(u_txt_path), 
                               's_analysis':self.data_analysis(s_txt_path),  }

        # 完善并显示报告
        self.report( model_id=_id, pngs_path_dict=pngs_path_dict, data_analysis_dict=data_analysis_dict )

    # 报告
    def report(self, model_id, pngs_path_dict, data_analysis_dict ):
        curr_path = os.getcwd() #当前路径
        #！！！注意切换路径！！！
        os.chdir( self.master.temp_path+"/model_"+str(model_id) ) #切换路径

        word_name = 'model_'+str(model_id)+'_report'#文档名称

        def ADD_HEADDINGS(document, contents, level, pt, CENTER=False, BOLD=False):
            # level级标题 
            heading = document.add_heading('', level)
            if CENTER: #标题居中
                heading.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            #添加标题
            heading_run = heading.add_run( contents )
            #设置标题
            heading_run.font.size = docx.shared.Pt(pt) #标题字体大小
            heading_run.font.name = u'Times New Roman' #标题字体类型
            r = heading_run._element
            r.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), u'黑体')
            heading_run.font.color.rgb = docx.shared.RGBColor(0,0,0) #黑色
            heading_run.bold = BOLD #加粗

        def ADD_PARAPRAPH(document, contents):
            paragraph = document.add_paragraph('')
            paragraph.paragraph_format.line_spacing = 1.25
            paragraph.paragraph_format.space_after = 0 #段后间距
            paragraph.paragraph_format.space_before = 0 #段前间距
            paragraph = paragraph.add_run( contents )
            paragraph.font.size = docx.shared.Pt(14) #标题字体大小
            paragraph.font.name = u'Times New Roman' #标题字体类型
            r = paragraph._element
            r.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), u'宋体')
            paragraph.font.color.rgb = docx.shared.RGBColor(0,0,0) #黑色
            paragraph.bold = False #加粗

        def ADD_PICTURE(document, pic_path, pic_name):
            paragraph = document.add_paragraph('')
            paragraph.paragraph_format.line_spacing = 1.25
            paragraph.paragraph_format.space_after = 0 #段后间距
            paragraph.paragraph_format.space_before = 0 #段前间距
            paragraph.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run('')
            run.add_picture( pic_path, width=docx.shared.Inches(7.5), height=docx.shared.Inches(7) )

            paragraph = document.add_paragraph('')
            paragraph.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.line_spacing = 1.25
            paragraph.paragraph_format.space_after = 0 #段后间距
            paragraph.paragraph_format.space_before = 0 #段前间距
            paragraph = paragraph.add_run( pic_name )
            paragraph.font.size = docx.shared.Pt(14) #标题字体大小
            paragraph.font.name = u'Times New Roman' #标题字体类型
            r = paragraph._element
            r.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), u'宋体')
            paragraph.font.color.rgb = docx.shared.RGBColor(0,0,0) #黑色
            paragraph.bold = False #加粗
            
        def ADD_TABLE(document, Dict, table_name):
            #简要说明
            name_and_number = Dict.pop('name_and_number')
            paragraph = document.add_paragraph('')
            paragraph.paragraph_format.line_spacing = 1.25
            paragraph.paragraph_format.space_after = 0 #段后间距
            paragraph.paragraph_format.space_before = 0 #段前间距
            paragraph.paragraph_format.first_line_indent = docx.shared.Inches(2)

            if list(name_and_number.keys())[0] == 'nodeLabel':
                contents = u'仿真模型共有节点：'+str(name_and_number['nodeLabel'])#+u'个节点，对X、Y、Z三个方向的位移和总位移大小分析如下表。')
            elif list(name_and_number.keys())[0] == 'elementLabel':
                contents = u'仿真模型共有单元：'+str(name_and_number['elementLabel'])#+u'个单元，对应力、压力、屈服分析如下表。')
                
            paragraph = paragraph.add_run( contents )
            paragraph.font.size = docx.shared.Pt(14) #标题字体大小
            paragraph.font.name = u'Times New Roman' #标题字体类型
            r = paragraph._element
            r.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), u'宋体')
            paragraph.font.color.rgb = docx.shared.RGBColor(0,0,0) #黑色
            paragraph.bold = False #加粗
            
            #加表头
            paragraph = document.add_paragraph('')
            paragraph.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.line_spacing = 1.25
            paragraph.paragraph_format.space_after = 0 #段后间距
            paragraph.paragraph_format.space_before = 0 #段前间距
            paragraph = paragraph.add_run( table_name )
            paragraph.font.size = docx.shared.Pt(14) #标题字体大小
            paragraph.font.name = u'Times New Roman' #标题字体类型
            r = paragraph._element
            r.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), u'宋体')
            paragraph.font.color.rgb = docx.shared.RGBColor(0,0,0) #黑色
            paragraph.bold = False #加粗

            #加表
            col_n = len( Dict )+1 # 列数
            row_n = len( Dict[list(Dict.keys())[0]] )+1 # 行数 
            col_names = list( Dict.keys() )
            row_names = list( Dict[col_names[0]].keys() )
            table = document.add_table(rows=row_n,cols=col_n,style='Table Grid')
            for i in range(1, row_n): #i行 
                table.cell(i,0).text = row_names[i-1]
            for j in range(1, col_n): #j列
                table.cell(0,j).text = col_names[j-1]

            for i in range(1, row_n): #i行
                for j in range(1, col_n): #j列
                    number = str(Dict[col_names[j-1]][row_names[i-1]]).split('e')
                    number[0] = str(round( float(number[0]), 4 ))
                    table.cell(i,j).text = 'e'.join(number)



        
        document = docx.Document( word_name+'.docx' )
        ##############################################################################
        #二级标题  应力应变分析
        ADD_HEADDINGS( document, u'3. 应力应变分析', 1, 14, BOLD=True )
        ADD_HEADDINGS( document, u'3.1 应力应变云图', 1, 14 )
        ADD_PICTURE( document, pngs_path_dict['s_cloud_png'], u'图1 应力应变云图' )
        # - - - - -  - - - -- - - - - -- - - - - - - -- - - - - - - -- - - - - - - - 
        ADD_HEADDINGS( document, u'3.2 应力应变基本分析', 1, 14 )
        ADD_TABLE( document, data_analysis_dict['s_analysis'], u'表1 应力应变分析结果' )
        # - - - - -  - - - -- - - - - -- - - - - - - -- - - - - - - -- - - - - - - - 
        ADD_HEADDINGS( document, u'3.3 基于应力应变的可靠性分析', 1, 14 )
        ADD_PICTURE( document, pngs_path_dict['pdf_mise'], u'图2 基于应力值(Mises)的概率密度函数' )
        ADD_PICTURE( document, pngs_path_dict['relia_mise'], u'图3 基于应力值(Mises)的可靠度函数' )
        ADD_PICTURE( document, pngs_path_dict['pdf_press'], u'图4 基于压力值(Press)的概率密度函数' )
        ADD_PICTURE( document, pngs_path_dict['relia_press'], u'图5 基于压力值(Press)的可靠度函数' )
        ADD_PICTURE( document, pngs_path_dict['pdf_tresca'], u'图6 基于屈服值(Tresca)的概率密度函数' )
        ADD_PICTURE( document, pngs_path_dict['relia_tresca'], u'图7 基于屈服值(Tresca)的可靠度函数' )
        ##############################################################################
        #二级标题  位移变形
        ADD_HEADDINGS( document, u'4. 位移变形分析', 1, 14, BOLD=True )
        ADD_HEADDINGS( document, u'4.1 位移变形云图', 1, 14 )
        ADD_PICTURE( document, pngs_path_dict['u_cloud_png'], u'图8 位移变形云图' )        
        # - - - - -  - - - -- - - - - -- - - - - - - -- - - - - - - -- - - - - - - - 
        ADD_HEADDINGS( document, u'4.2 位移变形基本分析', 1, 14 )
        ADD_TABLE( document, data_analysis_dict['u_analysis'], u'表2 位移变形分析结果' )
        # - - - - -  - - - -- - - - - -- - - - - - - -- - - - - - - -- - - - - - - - 
        ADD_HEADDINGS( document, u'3.3 基于位移变形的可靠性分析', 1, 14 )
        ADD_PICTURE( document, pngs_path_dict['pdf_x'], u'图9 基于X轴位移量的概率密度函数' )
        ADD_PICTURE( document, pngs_path_dict['relia_x'], u'图10 基于X轴位移量的可靠度函数' )
        ADD_PICTURE( document, pngs_path_dict['pdf_y'], u'图11 基于Y轴位移量的概率密度函数' )
        ADD_PICTURE( document, pngs_path_dict['relia_y'], u'图12 基于Y轴位移量的可靠度函数' )
        ADD_PICTURE( document, pngs_path_dict['pdf_z'], u'图13 基于Z轴位移量的概率密度函数' )
        ADD_PICTURE( document, pngs_path_dict['relia_z'], u'图14 基于Z轴位移量的可靠度函数' )
        ADD_PICTURE( document, pngs_path_dict['pdf_mag'], u'图9 基于总位移量的概率密度函数' )
        ADD_PICTURE( document, pngs_path_dict['relia_mag'], u'图10 基于总位移量的可靠度函数' )
        ##############################################################################
        document.save(word_name+'.docx')  #保存文档
        webbrowser.open( word_name+'.docx' )#打开文档
        os.chdir( curr_path ) #返回到当前路径

    def data_analysis(self, data_path):
        import pandas as pd
        from collections import defaultdict
        #import scipy.stats as st
        
        result = defaultdict(defaultdict)
        data = pd.read_csv(data_path,sep=',') #加载papa.txt,分隔符为','
        for col in list(data.columns.values):
            record = defaultdict(int)
            if col == 'nodeLabel' or col == 'elementLabel':
                record[col] = max( data[col] )
                result['name_and_number'] = record
                continue
            else:
                record['max'] = data[col].max()
                record['min'] = data[col].min()
                record['mean'] = data[col].mean()
                record['std'] = data[col].std()
                '''
                # 分布拟合，寻找最优拟合
                distributions = [st.alpha,st.anglit,st.arcsine,st.beta,st.betaprime,st.bradford,st.burr,st.cauchy,st.chi,st.chi2,st.cosine,
                                st.dgamma,st.dweibull,st.erlang,st.expon,st.exponnorm,st.exponweib,st.exponpow,st.f,st.fatiguelife,st.fisk,
                                st.foldcauchy,st.foldnorm,st.frechet_r,st.frechet_l,st.genlogistic,st.genpareto,st.gennorm,st.genexpon,
                                st.genextreme,st.gausshyper,st.gamma,st.gengamma,st.genhalflogistic,st.gilbrat,st.gompertz,st.gumbel_r,
                                st.gumbel_l,st.halfcauchy,st.halflogistic,st.halfnorm,st.halfgennorm,st.hypsecant,st.invgamma,st.invgauss,
                                st.invweibull,st.johnsonsb,st.johnsonsu,st.ksone,st.kstwobign,st.laplace,st.levy,st.levy_l,st.levy_stable,
                                st.logistic,st.loggamma,st.loglaplace,st.lognorm,st.lomax,st.maxwell,st.mielke,st.nakagami,st.ncx2,st.ncf,
                                st.nct,st.norm,st.pareto,st.pearson3,st.powerlaw,st.powerlognorm,st.powernorm,st.rdist,st.reciprocal,
                                st.rayleigh,st.rice,st.recipinvgauss,st.semicircular,st.t,st.triang,st.truncexpon,st.truncnorm,st.tukeylambda,
                                st.uniform,st.vonmises,st.vonmises_line,st.wald,st.weibull_min,st.weibull_max,st.wrapcauchy
                                ]
                mles = []
                for distribution in distributions:
                    pars = distribution.fit(data[col])
                    mle = distribution.nnlf(pars, data[col])
                    mles.append(mle)
                best_fit = sorted(zip(distributions, mles), key=lambda d: d[1])[0]
                record['best_distri_name'] = best_fit[0].name # Best fit reached from name 
                record['best_distri_MLE'] = best_fit[1] # MLE value
                '''
                
                result[col.lstrip()] = record
                
        return result


if __name__ == '__main__':
    app = QtWidgets.QApplication(sys.argv)

    m = MainApp()
    l = LoginPage(parent=m)
    l.show()



    sys.exit(app.exec_())
